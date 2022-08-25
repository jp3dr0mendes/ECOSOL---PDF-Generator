from tkinter import *
from tkinter import ttk
from tkinter import Text
from turtle import back
import webbrowser
import base64
from PIL import ImageTk, Image
from docx import Document
from docx.shared import Cm,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import aspose.words as aw
import os
import io
import PyPDF2 as ppdf
from datetime import date

#class File_Functions(File):
class File():
    def mm2p(mm: int):
        return mm/0.352777
    
    def base64_images_decode(self,image):
        return io.BytesIO(base64.b64decode(image))

    def set_document_margins(self,document):
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.right_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)

    def set_document_font(self,document, font_style:str):
        style = document.styles['Normal']
        font = style.font
        font.name = font_style
        font.size = Pt(12)
    
    def month_name(self, month: int):
        if month == 1:
            return 'Janeiro'
        if month == 2:
            return 'Fevereito'
        if month == 3:
            return 'Março'
        if month == 4:
            return 'Abril'
        if month == 5:
            return 'Maio'
        if month == 6:
            return 'Junho'
        if month == 7:
            return 'Julho'
        if month == 8:
            return 'Agosto'
        if month == 9:
            return 'Setembro'
        if month == 10:
            return 'Outubro'
        if month == 11:
            return 'Novembro'
        if month == 12:
            return 'Dezembro'
#Creating a class for documents notations
class Relatorio(File):
    def __init__(self):
        self.document = Document()
        self.project_atributes = {
            'Cliente':None,
            'Usos Finais':None,
            'Trocas':None,
            'Energia Economizada':None,
            'Trocas':None,
            'Demanda na Ponta':None,
            'RCB':None,
            'Lat Local': None,
            'Long Local': None,
        }
        self.project_client = {
            'Nome Cliente': None,
            'Nome Fantasia': None,
            'CNPJ': None,
            'Num Cliente': None,
            'Modalidade da Tarifa': None,
            'Classe/Subclasse': None,
            'Endereco': None,
            'Cidade': None,
            'Estado': None,
            'Telefone': None,
            'E-mail': None,
            'Contato': None,
            'Ramo de Atividade': None, 
        }
        self.document_images = {
            'Logo': 'iVBORw0KGgoAAAANSUhEUgAAAMgAAABGCAYAAACJ4ts2AAA1ZUlEQVR4nO19eXwV1dn/9zkzc/fsZCFAANlq2BsWlSVhUSkFtbY31t1qQa3aam1r7dufN9cu2rfV1tpaRa1V65arr1WroigQcQERRDS4gyyyhZD9bjPnPL8/7kwIIQlJiH1f23w/n/vJzZ2Zs815znn2Q+hDXDUN6dkCukuHdrAR0Vs2o6Uvy+9HP76UYAYBwHvf8zyx7zr/7r0/8R146QJPCABWhaD/77auH/3oPfp08ub5RF5utjYQAPY3mX4AKOvLCvrRj38x+pRALIapkswAYEpYfVl2P/rxv4G+ZX8YJCjFbok+Lbgf/fjfQZfzmBnkyBfdKozAvWgDMYMY3a+nH/34V6FTAgmFIIjARGAOfaEbAhOBCWD0E0k//o+hw4nPldDCYaiqCz3nVl/m/TWFofpaG8UAhVKE59p4mffBBxa7hwPg0BdLjP3oR49wxGTkEASVQ97xdfdxkwq1PxUP1q5/67u+RXPCsPqUSEIpInz3cs9vJ4/Rz5t9nLgTgKhAinj6rJ5+9OMYcBiBMECoAAdLkHHaGO0f6X6RCQk5Kg8PPHeBe8TcMKy+YLdWhaBTGFbVhe6Ligca35f1nBiap5/y7uWe2ygMhWD/LtKP/xs4bEcgpASB03LBzIiDgKQJle7Xsqbk4dERIzETgMkA2TJDj1EZhDYnDCvyTWPypEH6HQKsTAWhEeDTqRkAEAQQOcae/WtBwWDwCKKORCIMQPVxXSIYDB6xw0YiEYVevpPOwAAhCIHgod8qqsHhcJ/3CQCoMgiRWwwqGwtGNQhbwIhA9XaufSFwdojfnYIhtT/17eNfBNj8eSDBN6Xx1qu899j3HEZYjqZr77W+lfzLNOZfpvHbS72/bn8v2xqr04cic++PfB/wr9I4aZe96xrvP+2yBL48LBYFg0GtqxtCoZA42j3dQTAY1EKhUJc7q13PMe++HILoilMgSsmpfcEKh0IQXIkux4dDEJXBru/pDNTm0xscIVNQGIqD0CiCnWNzcMGckXjOMCBknM3h+dolby9xr6dw4i622aSeVLY6BG0Owfr4Ks89+TnaGKuFk4aHXHUN8rMHNscuZAZVVAD4168YAj1c6YPBoBaJRGQkEpEA9JKSkq8OGDBgsGVZmUIICeDgrl27PgiHwx8DKUIJh8OMnveNQqEQhcNhCQCFhYVDxo4de7yu6wPi8bjfMIzGaDS699VXX307Eok0tm1bD+sBYMug9g6xeBzyQye5Jw7wUa5UGhGpxqc/tD67erm5mcohAcCeK72rKwiNwpDh1L+ep8/2TJ6Yh6KkEi63ruTnjVxzw6rEOxTG/jZt69EY2jf2+P066JSwHALYfKnn+vFFxq+tBJu6BtFiwnp9pyw75f7Y2sogtPIIJDOICLzvR76X8zK1uQCwaYd10+RlsZ855XApdKqCtfYSz/XTj9N/bSVgCgFhKrZe/UzNnv9A7M3eDLYzGRYuXHiNz+e7RCkVA6Chm2pjIjKFEIG6urr7Xnrppd/ZE7nLwXTuyc/Pzxs/fvwSr9d7lqZpY91utyBKVamUQjKZbLIsq6qhoeH2NWvWvGg/3pOX1XrvjBkzyjIzM6/Udb3MMIwcXT+0tiWTSUgptyUSiWf27t37p02bNn3cG4J0xv+5bxsTxhYaP8700Hy/TgWaxx5Gi1Ef5YQp+ZOdDXj0T29G77pvE2p68d6IQyAKQ/31665Rs47TL8/x0hm6huFpAdH61qwYI2rxnro4v7K7Ud1+0r3x14DUrtMVm+fMx6Zr/XnCzS/phvA0RdW+y5+ILopsRUNPRIROtVIUhsWV0Kg8ftNnV4mpQwu1b5gtnPR7yD05Tzxy+xlpU4KVTQdDFRCwK2tbI9OhLbqyEhqVw3r6LPcp4wq1XyoTkgAInbSPdppXzH8g8WZvdiQA2L9/PwEAEY32er1jk8mkPUjdmxfMDLfbDSHECADYsmVLV0TlrOZq9uzZ5dnZ2b/xer3DlFJQSkmllGRmp9+WYRgBr9e7yOv1Llq4cOHd69ev/2FNTU0zukckAoDKyMjIPOmkk27x+/0XG4YBy7IAwLIsy2FFWdM0pev6cL/f/32v13v+gAEDwuFw+DZmJkpR7FEHw16d5drveq84Pl/cku4XbiQYIJhIsFAABEGmu8gQBo3NzaZf3JLjX3LuBHUtPRB7nBmCqFuET5yaoeqtpd7vjxxAN2YEtAyYDDAsJBgKECCwrkGmazQwPSDOyg+Isz7+gffeqx6M/TAcRmPbna4zBITSlS7GC6+AK6kGNCR7zqZ1ya9WVKeMhI+8Hb1kX638xPCQKxlnc0C2GHb6MOtvROCKLYdYPILdAAKYkQSAakCcdRbkbQu8g08cqj3gdwkhLUjNS8Zn+6x7J9zVO3btiI4IkTRNU0kpEwCg6zo0TevWx16Nj6rCdohj/vz5oYKCgsfcbvewpE2Ruq5rUkrNsqyEZVkmERm6rpNpmoqZraysrCXTpk17aujQoZmcot5OCdGWNdTkyZMLZ82a9UJWVtbFACzTNJXTVsuylGVZccuySAihCyGQTCZNXdezcnNz/7BgwYI7iQihUOioLLjN6qgNS7yXTh+q/yndIJcVYwtuAgCjMcFoTrAAwSW8JCABK8qJLJ8oKhuhRTYs9VxHBNUNDSdxMGWA/vQq319Khum3ZbhEhoxxEhoBBvSoBb3FZNmcYAEBF1wEmWBpAHLkQP2Se7/jW/GrucgXN3ajPoNYKkRhslJMMaQfpXUdoMtJEQ5DjQ1Cu/5V1I0faH67zE2veN3ksaJsDsnTF2253HMD/SV+I1fABSCZtPgTCJQm48xxS1UBwFhAMcNzxlfooQGZWr4Z5aThI9f+WrXx7CdjV3ElNAQhEe554zuAICKjqalpczKZfFUIoduraOsKysyHTRZmVoZhGPF4/BmgVfN0BILBoBYOh+XJJ5/885ycnAoppVRKkWEYrmg0ujkejz9cX1+/vqGhoUbXdVFQUDDS5XKd4fF4vu12u/V4PJ7IzMycW1xc/BARnY5DO0j7+ggAxowZkzZ48OB/pKWlTU0kEgnDMNzxeNxsaWl5zLKsp/fs2fOJZVlxt9udkZ+fX2wYRtDn851MREgmk8ns7OxLFyxYEA+Hw1d3JZNwCAJhqAdPcxWPzhW3QbFUADGg7zwg/769QT3/4ofyQ7eHjdLB+pjBGWJywIVzBmRpuVaMpRDApCH6zRsv9TZROHZHZSW08vKO2S2uTNnYPrrS88fjBmmXyRZOCg0uzYBr90H5RszkypWfqc3b6tSBDIMDXx/jKk734ZQBHvqmz0fCbOZEYa427cJJvmc+/Tw6D0BLN9gl0ebTYxx11SyPQKZWeHPD2ovFVdOHG/eygFRJlsOy9YoXzjGeFjeamwBgW516fHAeLmmJ8/afvZx4zd4GrXcud/+8qECbbbVw0nBDb25R9VUfWmet24UYqiGo/NiFcntVhqZpIhqNPrR69er/7mVRR2zbzgSbO3fuGVlZWb+QUpoAdKWUrKuru/6FF174I4B422eqq6vfAfDErFmzbs/MzLzT7/dPTiQSiczMzIXz58//0UsvvXRzRxM3GAyKcDgsFyxY8Mf09PSp8Xg8YRiGu6Wl5c36+vrL16xZs7GDNr8G4O7Zs2eX5+Tk/NHlcuUnEolkZmbmD0455ZT1kUjkoU6JZCyIALV9uPHTQEC4ZZytuGLzrV3qgrK/xdop2621AO7/0VTcfOl033XDcsTVICSFRlqmGyMAtNUIH4YUuw654TLP+aMKjKtklJOaAVdjnBu27LeuPvGexN/aP/PTVdbrAO5Zcb7nxAkDxZ15GdoEs4UTg/K0qT+d7/0ThWMXchAaeqkk6A66ZRmnlIFQp3Dirx9foU0ZOVi7XEY56g2Qb0SuHmQ2N3EI4sR7E1UvFurRqIVXqrYjjgqIBQ/BPShNu0CZUCTAzBAf7FMXlT+T+GRV6bGzVm3QSmSGYWjBYFBramrS09LSjlp+cXExA0AnwjlVVlaqMWPGDEhLS/sTESlmFlJKs76+/lsrV658hogwe/ZsPS8vj50dyLFVRCKRN4uKiuZNnDjxeZ/PN11KyYFAIDx9+vT7IpHIPrSRR5xJPHv27IXp6ekXJZPJpGEY7ubm5jWvv/76ooMHDzYGg0Ft//79VFVVpQAgFAph9erVoqysTIXD4crJkyd/MHTo0GddLlchM0ufz3drSUnJysrKyr1EdJjsYwuz8v4z0nIyvWqxSjJrLuif7JQ3lP0tHuG7YKx+GLy6LPVM2WqIsivAVI69v1sfveaNSzwbpg/THzxQJz+Z/T+xnzOD0IEcEgpBIAj13wt9BSMycSskS02HXh9XB1/6MLkw+IS1zmaXRGQLuLo49S7LAFE2Fkzl8TeuKcac6xb6nsnP1E6SUU4el6ddsPo7rkfpvuTzx6JJOxq67zoShmSGmDIl+oOnS32TC/O1E5BklWaIBQBuwFhg7S7E9rbwCstULwMpQWzled5J6W4xWEgW8JP70x3Wr6beE3uqL+SOzsDMHIlEZGlpKS1fvvyYBi4YDAoikgsWLLjG5/MNSiaTpqZpRl1d3ZJVq1Y9s3TpUmPZsmVWVVXVYX2JRFKLb2lpqV5VVVWXm5tbXlRUtFHTtKx4PP5fHo+nAe2EddvYp2VkZFQIIVgppcfj8R3V1dXBgwcPNpaWluqRSOSwesLhMACoqqoq2G3Z7PP5zi4oKFhFROzz+fLy8vKuIaKfBINB4bQLACLlEADkAE/yqwG3kSkIaGnh2kc/iS/jEETFpZBhQKHKrsv+zgDhLuh0afzvby3xFIDp0127EEM5NMKRE7ViLIgI6sOr1E8yMowBMsamIog3tlnnB5+w1vFSGBSGiXa7d9j+f1Up9DlVOJiTE/3WFTP86zN9VABBfHyO8Qsg+SIqob4oy1m3+TICGBXAhg0wn66Onr1jr3wNDJGXTl997WLjeCqHZICeejd568rtqHKMh8cNwGIjWxiNSa6r3mb9euSfYj/nypT++4vpUp+CIpGIHDduXL7b7b5MSildLpfR0tLyyKpVq+4rLS3Vly1bZqILHriqqsoKBoPahg0bdhw8ePCiffv2LVy+fPnvqqqq4mgzIWwjH8+ZM2ee1+udapqmAiAaGhqu2bZt2z6b0LpcUJYtW2aWlpbqr7322qstLS13GIZhmKapXC7XxRMmTMizWazWqRQsTn0fnSNGCx0MASQUtt38MmpRAQ53omkjgOlSmByCmHJ3/HdT7ok9CYA6WsUZICqH/NXMQG6BX1ykEqw0Dxk7Dsq/LXwk+dxbS2HQMphd9WtOFay7lsL4+Rrs+fSA/CEDmkqwHJBGJa9f4juZCMzdMCTKXpBRjwQXCkOFQhCXv4DPhv4+Ovv1bTJYF1fVmW79IgCESogfr7JeufK5+HZNgO+7EB6lcNrWnfLBZeutknF/if2XIIDK0V23iF6tC+0F8d6itLRUA4BBgwZ9y+v1ZiulKB6Px2pra28AQGVlZd2yZzgTs6qq6p9r1qx5oSvLutfrvcDWrIlYLPbGqlWr/icUComqqqpuLShVVVWKmam2tvamRCJRC4A8Hk9OXl7eaW371BZpHuGCIAIIUnKyu4NH4ZQmyV4MO3yfq0OpibtorDojPaBlCQAtzSq6Zof6JTOoZGD3FspLl6X8AKfcE4/sa1AbhAua0AlDMnAhgM6FHxvMTGnd7FdbiNTDdtBSNz4VY0GrQtAFQc24L/Z49q+j06p2JJdXBlMaCif4STHw1hb4X/zYvHjErS0X/HhFYttbS2HIG6Azd7tOiF5MdXKsdccIhwBcLtcZAFjXdZFMJp9cu3btJ7Yw3RPrLIdCIREKhUQHwjJFIhE5cuTIdJfLNUdKCQCUSCTuAYDVq1e32pq6AVVeXi7Wr1+/N5FIPKlpGhERezye09v2qS3q4yoJBYZiuHUayoAHEYjuhB5QGIq6CJQrq0jtQvlpWAwCQydR28LPfeepxGeIHN2W0Qa8OjUduDaGu0BEMJnTXJh/zQnIdjiYbpbVbegA0FUHO0DKxYAhNlwKbcoyRL/3rHzJudi2rD+vRy1g1XIltA0vQUxxttK+Uem2B3XyvbcQ4XBYTZ8+PV/X9SlSSiIiJBKJh3tbfmcEZcsGcvDgwZN1XS8EgEQi0fz555+vAFK7Qi+qo0Qi8Q+fz/ddIiJd16eOHDkyPRwON9rt58iW1LvaUS8+GD2AiRlWeoAGbb7MczWVx29mhihbDb2sCrI3DoOMlOzxm5OQ5tUxDSYTa0BzXEUYoNXVPRvH1YBiAC9v5eeHZsqWgEv4A24x4Jyx/pLfr21ZsToEDX0s1+oA8Pol3lvyM0WxSiopUp1iR4YgQsqqRWAwNBAaqvequ4gSqwEoIuCGGzo3/Yfs+BIA8uR8+G8Ner4XcGuzZVKBqNXAiLb1AYACKTbgTcTV38f9OX7/0TQVbTcNZqZQKCS2bNkijubg56C9W0YwGKRIJIK0tLRJmqZlEhFM06zZtWvXWgBsC9R9AscbwOVyjdd1HcwMy7Leq66u3mnbcXpUl+PZu3fv3rfS0tIOGoaRLYTIHThw4LhPPvnkddvgycFI6p29W9Py9pTBvvp0n5YhTZYj8/RfbljiUUTxW4DUhONKaJEIEOyBd20kCIEI5Jgib7FLE7kgoCXKsVf28ptjAeYeyqHhMFRK8xbbFRzrezfgpRM0AWR41WQAK8p6Ulg3oQPA0Awxt3CQNgkx0bVUYtt/89P42x99X3vs/b3J20+vtN4Md061VBEG538DeWUDfafnB3B1doYoTvFNR5m3CkCAoGrNjQCwISulcelmvxL2ap3s5v1HwJm0hmGMdPyeLMv68IMPPqhtb3w8VuTl5TEACCEm2H9BRBsBoKysTAN6vCoyALz33nv7RowY8RkRZRuGITIyMoYBeN1m2RQBbNsnDi4ajUeysnB5shnKrcP4apH+m/0/9n9rW4O17LkPE/+kcuxtLbwSWnfc3h0lQMlAMcrjJQHJSCretdKV2An0chuugEaApRM2QsMJUICuYQIAYGzfO7nqAGApblBRTlpxVkSdz1y7Q+wz4B41XD/Lp/PcOxbgtMuft9aBIKid1oNDIBGGesrj/lZRNv3Zn6lBNSopFctueItZmoCelIgBQMnA7nXetheOmjFjxlgAuqZph7VJKXXEeyEiPZFIfPLmm2+2sh/OpJVSjrLvAYBqACgvL+8JsXYbXq+3wOmDaZo7jqUsx6HSNM2tXq/3qwCQSCSGtr+vojzFLfx+AX4xwKsWZ2eJwWaUTaUgcjPE1NwMY+qIbH3vdyarFz88IB8/9e/JVVSOZqD7nryNcS4EAdAITXH6NJJycO2u79ZhWI3UvDkYx85cmbKj5/sxCABQ/QURCMAZwkMul8LR9VoCaGxRLQe2m3ds3sO3X/68tQsA2hMHYGs5AKJHEndUBvnlyYPUdfl+7fw0H7m67AoBkHDBQ1BN8HezL2SzQTIjI2NJIBC4uINSHQHeIRIGIIUQem1tbTmASGlpqdZWnWoYRl6b7/uBQ7tLX6GyslLZBJgOpAgkkUjsP5YynV3C7XZ/7vDMmqYVAod2LCBla6iogPjhC9E9hT7f6fNGq8cHZInhKsZsJjhJgMjxUUFOun7B0CztgrrrjI8+b1aPr93GyygS3y4IkNy1u4cpVQFYAwjwueydKNK7DaTMlps+b1Y7xuQKQANiFqUDgHYjFPpYUNcB4J296q49LeYoJZEEHVkBMyil7gJJXSW2HlCPXfCk+R4AhB6DqAilbuuogkgQIhQBl0eSHwK4+NFv6MsG5xinawqC2wj0jgzi1MMCigW7E0m12r6lW6sNETERKU3TOhuow8oRQigAlhCiQ6Opruv+VPsY8Xg82p029BAkhHB2K59TF4AG4PDJ3KvCieJtvns7vMdW11I4ujF0Ik44b5r35vyAdmFagFwwGZYJCZMtEPTMgBidmSF+NjRDXVYyxHPr5Lviv075pnZAJDbL49Kp1U3QKzg1hj0U0NvDLezoUwakYjcAnXvOih4VOgCc9mjizp4+yCHoG/aApiyDGe5CK1Vub8FvLYVRMh+Kyq21tk9Pj9ADdSCEEAI4XHDvAppz+1Hb0Idyx9Fgq3qPGd21CR0iEuwPvxG7+K+L9L/MGO66OMdHZ+T4qQAGaUgyrCSbDCDgpuxJw4xfbrtaTPvJi9FzEUSUKzrWiIo2iy73LnfaEWhLCXxsQYNdQgdSceK5xT2rwHETuWWud9DQHJ77rUj8IQKUSvnjgAC+bQHSJ+Z5zih7IB6ZsgwxWtbzusoA1UPi0BsbG1ckEoknhRA6EbFSynkpTnxEK5hZAdCj0ejrANDeIGdZVmuGel3XA91tRw/ASikiItY0rXWHIjq06h4LlFJu57sTCtAZyNYSpaq31gPW+ltOwX9NLvScfFym+EaGmxZmplMaFGAmYAnJPGygdtpN83x/o3A0yIfSGhyGuKQm53tCkq8v+kWKU7thqrUmAKtbgS89hA4cWuW7gzbRXO5Nl3kvLhogbqypVw8AeFBVQiNqtZNQpBwtI3LEZfv/n//63TUyPOnOeGV5BPJYkj50BWaGHRfxwqpVq/7S22La/mOaZg1gK/SVyj3mRnaANkJ/E5DSYnk8njyg9/KOw5pZllXo/KaU2ne0Mu0dIJUscCyIynEQiD8G4LHvz0DR+cWebw/N1q7MTRdDrARL2cLmiHztm69e4imncPyxVSHocxytps1GacB+Z0lsTqqUTFfRuxCH1fbimu1BkSMvpxncAIDlDa2Gxz7bTXqU58qJ4nruXH3qtMGuv+aka+PYZKyvt+4DUi4kD53pWuxh3iGE+Q4zZPXluLe4UNyTa9Ajn13ju/T3G6PfwGo0MPU5kbSWJYTwlJaW6nb/usWX2jtHaxnOJBJCfAKkQmgBFLe5t8+RTCb3ud1uEBFcLlcR0GsZhBxrvaZpx9lth8vl2tbtAuxduzWzSSWYCDv++Fr8vy+bgL9dW+q9bWS+/m0rwQoEHpklrgRQWZaa+Ict5nkB7EoJC0DARccBIJHKptnjCVBm/81PE0OgpVRDtTFb8B/7BVnSu4MQIFABvvZF5E0d5KrMydSGwWLrQJPacurfzQ8cf//ZQ7WlMZNeYzY3MwMPniZXDs2kZr+bPEMLtLIrx3vuJooHOQS9r62eDphZVVVVWaWlpTiag19ncCamUupjKSWICEKI4sLCwpzdu3fXohN2ojdwiNGyrHftOsHMk4FWo19PQQB43Lhx+ZqmDWdmSCk5Go1ubdu3bhbEiECmvoAQgkZh7L9zc+zsz3/oTSvM0b8uE6zSXDT16fM9Q4jiO5zkCo6lfvPnaussj2C3AfLqNOTmUu+gn1bFdh0ttrxD2DuPKTERMuWL1JxMqd+PVfDvCN11VqSKypTMe8U0//0DsrRhZgvHYJDeaPJLACxUgxeVwJfhprJMD05BavKIC55JbKtP8EYYpFnNnBg50PjWhks8P3FiTPqqI92NQe8unIlZU1Oz2bKsJjt2Pef444+fCnScB6u3cCZsMpl817IsMDN0XR83ceLEQbDHsSfl2bEolJ+fP1nX9WwAkFLWbNu2rbpt33oKApjCsFaVQhcErNll/VcsxpYmAK9BbrL4eACIpMKw4Vjq39oZr05KrgMDHg8FTh2DKQxQWQ/7Zbuu8K/mIj/dTZMhAWUykhKpILItfc+2d6uBHErtDpsudd84vEAssKJsGi64W5pUy+4mea/Del0/3jcrENACPoNK7voGBgqRkkd2N6l7ABAIQpksxwzUb3ruXPc8CsPqjptyT9FHzorMzPT222/vtixrk6ZpbMsG56CXO4fjrNj+d2fCbt26daNlWXsBwO12p+fl5c0HgNLS0t4QI7vd7jN1XYcQgqWUG7Zv315v139E++2ApW6N25wqSMmgb1ea1ZbF2yAgSAAD0zEQABwlDNli4fWvoq4lyRuhE2s6IccvziCAy3po+V4dgkYAFozwnOL3URoDaI5z04pP6U0AQGXfJ7Q76sA7UX8rznGdXjzQ+LlKsEUCYCLx4X753dn3JbfAZtUGZ+BMEoDfS+lT83wLmQFeCmPa3fEHt+6x7te9ZEgL0u8mccJg/YGbSzFYPA7ZF+lM+8qDty1sNw+YpvkkEZFlWcrtdp8xbdq04ZWVld0xqx7WxHA4rMLhsOqASDgYDGrbt2+vN01zpWMX8Xg8lwA9lkNEZWWlsmNYzrQsiwGQaZrPAK0GxMMrRyoFD7pP+AwABFiGgZTmjQCXOJLAVqey3qA+xs+AQCrJnOnBGX9a6CtAECrUgzEsG5vyCyxME0sIBNKJm5LqlR+vaNlvW+b/tTsIhyDmVsH6zXzXqKnD9HsNQawsKM1Lxqd7zd+X3B1/lEPQUZGSJbwujIadHwbgUW3LuSwSu2LfAbXJ8JLLSnAyK1MUnj3Z+yAzDIzt8/NB+qQsx4t2z549lbFYrJGI2O12p2VnZ4eJiO2JftS6nGCoE088cfYJJ5wwoys3+Wg0+pBSiizLkl6vd9bs2bNPs6Mju8WOlpSUaETEgwcPvtbj8eQwMycSifp9+/Y91bZPbUAE8AOne04CoHMIohvvgkTKOGgkkwiAACWB3U2iDjhk7QZSHrgA8Nx2PNHUrJrB4LSAyDh5uPoJEbhiafc4iFUh6FQOuepC94KcdDFLJlmCQTsb+EHgECF2BasvA6YYIFQAgwbDe+4E/eGMgJZjJtjUfeTaWyOrRv05/mOuhIaUR2ZKoOU2WxzZAvhAcGQLaMU+tLyyw/x2c4tq0F2km1E2i/L0sncv8/yGyiEROjZWq6+CpNpBBYNB7Z133vk8FovdbxiGZpqmmZaWdt7cuXOD4XDYWrp0qY4uJlQwGNQef/xxOXr06AF5eXmPFBYWVs2dO/fK4uJiV9vnIpGIZGZatWrVilgstskwDAFAZWVl3TJ06NDMV155xTpaCtOSkhJjw4YN5qxZs07y+/3fN03TcrlcIpFI/P3tt9/ebT/f+o7seA8OlaJgYbFY/u7lvusoDIUQtK5iQTgIoUIQT55jjDV0GsoSSFpgxfIjAEDxIQIJh6G4Etq1z8Y+r2lRjwgvCRVna0i2fsVT5Z4ZtAxmZQiurvpVGYQ270ZYV01D+oQC7U+GRkpzER1okB/fWh1/mhk05wuKUO18VQpBI4L1/hXe2wblalPMVJpQo6lJ7Xn2Yz5HEGRFNUTYdjMAALQ5jYrafC+PQKa0XMkPX79EWzK1iCqFBiXjbH2lQL9mzUWetygcf9jRhPWyL4fqJiJHzVtaWtqth+2kB0es7HYSBjpw4MBNdgqfHKQm7r0nnXRSw7Jly15kZiovL9eKi4vZjhFHMBik/fv3kx1D7h89evQjXq+3UCmlMjMzbzdN81kA29AmLt22h5hNTU0VXq/3H8xs+ny+kePHj//79u3bg5FIJFZaWqrn5eWxU1coFMKWLVsoGAyivLzcnDRp0qjs7OzHdF13MbOMxWL1+/bt+y1Sqt/DWJCKLSCUQv/uV32/y8nQ0gIu/uU7l3qjFI79Hkh5S0S2gIPF4AoAFQCwBYRgSqW/9Qf6tV4f6cpibknypx/XJj8lAkR7zVR1yiHyb2eoX+cH1Le9bvJ6CcbM47RHHzrTNbc8nPyYQ9CxBYxicEUYQAgYuwUULAbZRmnXdTO9D2enayPMeCpl7dZ6FY6sTcXC4+jOo1SUAZ1D0LEHxEd3flUUhuqQQOyJar16oXvJVwbqS2SMLU2HFrcYm3Zb53/36cRuJ6/qUSo51LpyJ31QLLLlcu9vjy/Sf2xG2TQEqfEDtbse/YaxmcrN97qTMa/DNrfZQSzLStjq3W6reKuqqjq7pOxsI3sCgcAPCgoKHpZSmpqmpRUUFPxj3rx5PyWiO9rX5SRHmDBhwriioqI709LSZiSTyYTb7XY3NDT8Zs2aNdvap+KJRCLS9sJ96mtf+9pDWVlZ5yYSiURaWtrXFy9e/PyOHTuuqKqqqm5bj0OQkUgEs2bN+npOTs4dHo9nsGVZCcMw3AcOHPj5hg0bdrSvK2RnUvzDfE+RR6ezkWClM3jCYO3Wj672Trp/A/+cwvGdh9XV2lBg46We64dkaefJOJuaj4y6mHziB8uR6CgZB4WheCy07zyV+GxsHt04dbjxWxnnRHZADF40xnh55bm4nMLJZ9FhZcCyr7lGLyjW7xyUI+ZYMY4bfvJs3y2fnn5P/KHuehQzg+9Yj9o71vfMtHAEa+Dk233yHF/JycPFq34XDNOCMnxkbNlu/mTsnfHfth8EJxfq3mt9K/OztDnAkbl5nfrs5GHYcY3vxSF52txkC5suLxl7auV7Zy+Pnri6Gi2Oq0p3OuAkM1i8ePGfA4HA9yzLUrFY7D0ArwkhWvNlHT5Y3NazVwohvI2Njc868d8d7STOBJs3b97N+fn51yWTSRMpd3qKRqPrW1pa7ne5XGv27NnTaBiG8Hg8QwKBwJler/cit9udbpqmQxwrn3322a8xs2nrFdq3TzAzT5o0KWPIkCEr09PTJzuJ4xKJRFMsFnsgGo0+kUwmdzY0NJgFBQUe0zTHp6enn+v1es/QNA1KqYTL5XIfPHjwgeXLl1/YWU4sZzFa913vZdOGan+BhLQUlO4ho65R1Rxo4cd21Mmno5J3vrNXxE4dI9OaW7SxI3PEkiE5Yp4yYQkdWkuc6598X004739inyNlAO5wgbMXXvXxlb7KkYO1b1ktnNANuBNJYE+zenpHHT9gCvl+ZKPWMKNIeYdkY/CgNO3MvHS6IMMvMswYJww/uQ/WyY/+uCY6q+J11FQQKNw+zMKej/wz30BTo22GTu5YkmMHWvhpQ6NEyoDFDrvKIGKwrXggSCL276zjR6fcE3vy8GMMAEIl1CUnInt6Af3D7yOPGeWkkUauXftkZUfE0UMwqsGCoB55N3r+khL/m1kBGmTGODkwTxt391zf3UTRs5lTuVl7XDizUkpZPp9vghN81I1n4Ha7EYvFGMD/dJabNxKJODvJT08++WQ9Ozv7WjvyL+nz+ab6/f6pyWSSi4qKmoQQmmEYfieXrmVZpsfjcTc1Na39+OOPzwKQbJ+jqg2Ufa3eNM0zRo4c+VQgEJhkpxVNy8rKuiItLe2KZDIZGzBgQELXdZ/L5XIJIWBZlqWU0lwul7uuru6R5cuXL+kqIvGQg2LszreWevzjCvXfuQ3SzCgnszyUm5UmrhyVK66sa2ZrWiG3+F1aRmAQpWZRgpOaBkMCtKVGfu/8J+O7zusk7U8rylO+XsOGRc9/5Rs+vahAO0NGWekCaliudtqwHD6tvlmo4vl80KXp/hw/eaETOMmQcU4aAXIfqFOfPPmutTj8BvZXVEC0J44j+ggoaUHpGnmGDBBnHVVMlwD8hIaE+TGAJw8TxAgAKkDTBiCesHgjGDA85Kqrk+8veSa6xE5TeUzCEIWhHv0WtOtewu5Nu+W58SQswwVWJkCs1gMAynulhXIZhiF0XXc58Q89+RBRl458sMNsQ6GQWLFixY9qamq+k0wm97hcLpdNKEoIwR6PJ93lcvmISJqmqTRNgxDCqK+vf/i9995b8NFHHx3A0ZNXKztkeMemTZvm1tfXP2Tn+oVpmszMyuVyeXw+X6ZhGIZSSkkp4XK5dGaO1tbW3vD888+fhxQhAl3sxs5xF1OWxW9Zu1WeebBZbTd85IJGkAm2lAUry0ciP02kB1wkpQlLmcyaj1wtFsc27VBLp90df7SyG6wOpaRU3rED8aF/iH5zyw7rpriC0rykqySzTMLK9BAPTBcDcnzkVhKWSrAkF0Fzkeuz/dazv13VMmfp88mPusWKG0y6IK/mIWG4iKAYkEf5KGZIhkuk5kN7GYQRBi4FogtGRs9/8EzfWwEPHffGTvOs5Z+gMbIFWnkXL1Z188joQ+lM41WbL/VeM36UfvuOnda9o/4cv9XWZ3ebCB0bQSwW20lEn0kpYw6BACl2ygmSasdaOdeteDzuSyQS3Yng43A4zDYb9rcxY8a8OHz48CvcbndQ07QRuq4Lm60jZtaklImWlpb1jY2Nt1VVVT1ul9Gt4w9sNk/s2LGjbseOHefNnj27Mi0t7QrDMGYYhuG3w3LBzCSlJNM0a1paWv5ZU1Pzx/Xr12+yi+mWOwy1KlFiT940E6vPnOy7LMOLc7LdYpyhw05JToBiTWOgPs4ttfXq2Zc/xW8u/Wd0Y0/eGQGMFAukxt4R+9mD3/A+fVIRXZPtpa9nesmfMlcSABZCQbTEGU0tas0HB9Tdc+6PPwgcYg2PVtfefcJKeuSHgsiwNawaAGIcycITgVkBRLBUM/l2N3GdM4BHwGnAivM9J3oFjpt5f/whRzbp8H6b59v9Q9+qgdlaGdCpDNJhPW9/13Pdmw3xO5dWorEn8kc7uAB4kRKWGa0R9IfGoJP+MlILRQw9iGFvy9fn5uYGRo0a9RVN0yb6/f4CKaWKxWKfJRKJDevXr/8IaCVUp76egOwdjgFg+vTpw4QQk/x+/0jDMLzxeLw+Ho9vaWpqeu+9997b16ZtPT6Srd079jxdbhSzLiaMytSG+Fzs3t3MDTsa1QecEJvOfiolwB9D2k9HHpUA8MjpniG6R311RKYoHuAlfywJ+UGt2tZiqfXnPGFWAynz2tGiF9tB5AD+Wvuf7BRxdPg8AezcNxTQDCD2CdA5V9HWWHS0/EhONODuH/pWdXUEW1fPfklx1OPViAhHu6c7CAaD2tFsPZ25svQEDFB3TjNm7vqYtm7Xl0o812U5glKEeKx19QZdDrgzAEfbznq7g7Q+XwkN5cd8WGOf+F/1tu5QKERt3Tja2Cqc3awvQEg5SlLbmA47aXafHuJpe+4StoBg+1atBlBjJ5fu64M8QyGIsVtAucWgMudH2y7S02PX2qK3HhrU+ngf4Fh2kH704/8y+nTyEkEqtiMK+/74436g9bi7f1ls/H86+pRABFGa8JEGAJpAhxk0+tFrEOxFp68T1/Wjc/RZ0A8ANCbU6uZ6tby5Xi1viqv3AXwhQSz/oeBp06YV5+bmBoiITzjhhP4F6F+AL7MGqUO01RjZJzF1qCBwTmpq+1sn91LbIwNsgfiojnF2ArrD4tw7gQgGg1RcXMyrV68WHTwjAKh58+ZNdblcTxHRXillnWVZd7z88stPdOZGYhsa23oL96kQ349ewFHZ2Z9/OfEdq4qzB+iTvnWzvQQAs2bNGj5t2rTi+fPnn7po0aJfFxcXB9pe70W5/egG/m12EMfJcPr06WW5ubnlzOyVUr5ZU1Nz/4YNG6JoZ1WeO3duMBAIjJBSKl3XKR6PN2/evPm+PXv2OPcCAE+ePLkwPz+/3O125wFI1tbWrn311VeXdyIHEACeOnVqTlZW1tmJROLRqqqqA+3rbntvSUnJ+IKCgiVElCulfHv//v1/3bBhQ2fPYN68eWdt2rTpudra2qauyp05c2ZJZmbmhUqpdMuyql588cX7jsFY+R+Lf4uVxiGOuXPnXlRQUPCSlHKolLJFCPH9tLS0yfY9h1nSPR7PTUqp65PJ5HnxeHyJaZrnaZrmca47bFV+fv5sj8fz+2g0mh+NRo/Pysp6csGCBTcSEbc3ADrPpKenf33QoEG3a5p2etvf27YXAM+YMeOEwYMHrwEwMZFIHCCic3JycqYAqTNDnPvteqisrGze4MGDHz3xxBN/iHasX5tyMWfOnNlZWVkrlFKGYRjvu1yuX9rnprM9Dv3oJv4dbBQUDofV+PHjs/x+/x/j8fi1L7zwwm32NW3o0KEG0Orb1NYHK9nQ0HDxmjVrnuigzNYVlohcpmluXLFixSUAUFZW9vX09PRHTzjhhD9EIpGDaLOKl5WVqaqqKrjd7rk1NTV7vF5vKYB7nd/btzsrK+sPpmn+47nnnruoze860Hps22Ft8vl8CxsaGrYx83Sg0/xcHAgEbrYs6/bnn38+BADTpk2rzM/P3zhz5sw7wuHw5s5c+vtxJL70O4iz0g4cOHAuETW/8MILfwyFQsKOKJTbt2+Pd/KobhgG2Y6FoqucD0Skl5aW6qWlpXo8Hl+paVptQUFBPtC6MwE2oU6bNi1dCFFSU1PzEyKaOG3atPR2xOnsdoVCiK/U1dXdzMy0dOlSw77eXlFAzum3mqaV7tu3LySEyJ85c+YoHJ4SSITDYTVx4sRBSqlB+/fvvzcYDGqlpaWeN998cxuADYFA4BSg48QN/egYX/qBaqOJmi6l/AD2am5rpDpzUAQAxGIxw84Gr3WVV4uZUVVVZVVVVVmJRCJLKeVpbGxsBFpPpmol1PT09JkAxLp16x4DoGdmZs5ue905P10IMZKIWph5OxHxsmXLOj2ECACXlpaOBzDgjTfeeJiZawKBwDzgUEqgUCgEACgqKhricrlaNmzYsMvehSyksprsVEqNAY49Y/x/Er70BOLAMIwCKeVBAGij3jzC6uzsFEKIloKCgttOO+20daeddtoHixcvXgYcijZ0JhEzJ4ko/bjjjisaPXr08EGDBv3Vsqz3Vq5cubttjqk2x6idrJR6G4AppdxgGMbX2l53YFnWANM0rbVr13Z5BLJDAF6v91Rm/hiAVEq9oev63LbtdJCKnRISh3syMID9mqblHHUg+3EY/m0IRCklYfv7238xa9ass0455ZQfAYfsI23iRERjY+MTe/bsua2mpuaW/fv3/wM4REAOmDnmdruHHX/88c+NGzfuHWbO2LFjx1Jmbo0HB4DVq1dLACSEmCmlfBYASSmXE9GJAMi+3rZcSUSipKSkS6HZOZVW07S5SqkXASAWi60AMH7w4MHe9mefdwZm7jJzSD86xr8NgUgptxmGMQgANzU1EQByu91jPB7PZKDDjOYu0zSfXr9+/cNvvPHGHevWrXvO/v2wFVnTNF8ymXwfwKlSyveVUlXvvvvuVjsDibNKCyLiWbNmDSOiAfX19SsBcF1d3UoiyiktLR1hq4RFcXExA4Cu65/rum4A8OHQBG/fRrLlikwiGh2Px18AUhkYiSj5la98ZRpwmByERCJhmqbpGzhwoA8ANTc3UygUEkKIQinlnk7Goh+d4EtPIM4Ky8yriWhcYWFhzvLlyxNExMwcE0I0d/QcM2vO6VFdCehSSh1A07PPPvt5XV3dlYZh/GDq1KmjnfBb4JBcEQgE5jPz5+vWrdsXDAa19evX7wWwze/3t8oLjsxy4MCBj5VShtfrHQeAQ6GQhnbE6cgtubm5JwFIrFy5sjoUCont27fHmbna5XKdCqSEbqfcrVu37lRKZRQVFY0EwIsWLeJwOKx0XR+jadrbvRzm/1h86QkkHA4rZqYVK1a8rpTaWlJScn9RUVGWHWueDnTsNElESkqZW1xcnD1p0qTCCRMm5HUU2ERELIQgAPTKK6+sl1I+W1BQcCsAbp/gQdO0rzHzqwCwf/9+AwAsy6oSQnwNaJUX2E5GV6+UejYrK+uWgQMH+sLhsAXAVVJS4mizWld6t9u9gJk3AJDr1q0zAEBK+aKmaXOBVvbOcS/ZS0TVubm5v7DHx5o3b95ZAIYeOHDgn8AXd3zDvyO+9AQCtMaYW/v27TtLCDF64sSJnyxcuPAVn893tWma2zt6RkqZyMjI+PPQoUM/Kigo+KCwsPDDffv2jbYvHzYupmlqSK3y4uDBg9cy84zZs2fPcVKCRiIRWVxcHJBSlsTj8eUA0NzcLAEgHo+/oJT66siRI9MdeSESiTAzU21t7dXMzFOmTHl/8eLFzy5YsOADRzsVDAYdXy4wc2kymXweAGpqahQANDY2vqyUGjRjxowix6Jvy0R08ODB7xHR5EWLFr29aNGi5T6f7954PH71+vXr93aWvLofHePfiRd1DHZpCxYsmO9yuTKbm5s/W7ly5ToAR7iazJgxY2IgEMh1JpeUkgGsW7FiRQvauJrMnDkzF8DAV199dbNTxoknnjjO5/NFX3755a3ObyeccIKXmSesW7duI4C2milt1qxZJdFo9D3b5aV9e12nnnrqNxOJxFf9fv+be/fuXdXe1WT27NlTW1paqtu5zNBJJ51UYlnWR22Pr3b+FhQU5E6ePPlkTdN8dXV1b7z22mvV/W7y/fhX7Ih9uag47dXnz5//5z6so9+BsR+dgoLBoOZ80MlkC4VCou19XSRWoPaTq7Os7p2V0UXZAimfqkmLFy9ev2jRIkejdVjZndV3lOzyrePQTxz9+FLCIZxTTz31lyeffPJPgP6Vvh/9OALFxcXZ/9tt6EfH+P8lNHY5bRPP/AAAAABJRU5ErkJggg==',
        }
        self.cnv = None
        self.date = date.today()
        self.set_document_margins(self.document)
        self.set_document_font(self.document,'Arial')

    def set_paragraphs(self, paragraph:str):
        self.document_paragraphs = {
            'Resumo Executivo':f'''
                \n \t O objetivo deste projeto é promover a eficientização da iluminação {self.project_atributes['Usos Finais']} do {self.project_atributes['Cliente']} e implementar um sistema de geração de energia fotovoltaica nas instalações da unidade.
                \tÉ prevista a substituição das lâmpadas fluorescentes tubulares e fluorescentes compactas por modelos mais eficientes de LED {self.project_atributes['Trocas']}. Pretende-se também realizar a instalação de um sistema de geração solar fotovoltaica conectada à rede. Essas medidas reduzirão consideravelmente o consumo de energia elétrica, e impactarão positivamente em todos os meios envolvidos, tanto na redução dos custos para o beneficiário, como na redução da demanda no horário de ponta para concessionária.
                \tA fim de maximizar o impacto social e a redução no consumo de energia elétrica, está previsto a organização de palestras, workshops, distribuição de material informativo (folders/cartilhas) e treinamentos, garantindo uma eficaz ação educativa, de forma a dar ciência das ações adotadas pela Concessionária, da importância do projeto de eficiência, e dos benefícios trazidos pelo consumo consciente.
                \tPara a metodologia de Medição e Verificação, serão utilizados os procedimentos do Protocolo Internacional de Medição e Verificação do Desempenho Energético conforme os procedimentos do Programa de Eficiência Energética – PROPEE e seu guia de M&V.
                \tOs resultados esperados com o projeto são {self.project_atributes['Energia Economizada']} MWh/ano de energia economizada, {self.project_atributes['Demanda na Ponta']}kW de demanda na ponta e uma relação custo-benefício (RCB) de {self.project_atributes['RCB']}''',
        
        'Apresentação do Cliente':'',
        'Apresentação da Empresa': '''
            \tA Ecosol é uma empresa concebida no início da década de 90, idealizada com o objetivo de propagar a cultura de eficiência energética e energia solar. Localizada em Niterói – Rio de Janeiro, a Ecosol vem desde 1993 oferecendo soluções em aquecimento solar, instalações hidráulicas, elétricas e Fotovoltaicas. Nascida na Região Oceânica como fabricante de Coletores, a Ecosol especializou-se em projetos e instalações, participando de grandes obras como o maracanã, as arenas olímpicas, indústrias, hospitais, escolas, academias, prédios residenciais e comerciais, além das diversas participações no Programa de Eficiência Energética da ANEEL, com projetos aprovados em todas as macrorregiões do Brasil.
            \tPara o sucesso de seus serviços, a Ecosol conta com um corpo profissional que disponibiliza as seguintes estruturas:

            - Equipe de projetos;
            - Equipe Técnica;
            - Equipe de suporte volante;
            - Equipe de retaguarda para suporte permanente;
            - Manutenção e reparo de equipamentos;

            \tAtualmente, presente em 18 estados do país e com mais de 600 sistemas de energia solar fotovoltaica em operação, a ECOSOL contribui na construção de um futuro mais sustentável para o planeta. Em seu período de atuação, pode-se comprovar experiência e credibilidade junto aos grupos de concessionárias de energia ENEL, Energisa, Equatorial e Neoenergia.''',
        'Vistoria na UC': f'''
            \tCom o objetivo de identificar as necessidades e a possibilidade de eficientização da UC, foi realizada, pelo engenheiro eletricista subcontratado Lucas Araújo Pereira, uma vistoria técnica e o levantamento de dados mostrado na tabela na seção “5.2. Levantamento”. Como pode ser observado, a proposta de eficientização contempla a substituição das lâmpadas de tecnologia antiga por novas lâmpadas mais eficientes.
            \tCom as coordenadas geográficas, foi possível verificar o potencial de instalação do sistema Fotovoltaico. As coordenadas encontradas de -4.9654 W, -42.7958 S , foram então inseridas no banco de dados do CRESESB, que resultou no gráfico abaixo. Os resultados obtidos foram posteriormente avaliados para dimensionamento. Foi considerado o sistema fotovoltaico, pois além da representativa economia de energia, os outros usos fins demonstram-se suficientemente eficientes. Segue, levantamento fotográfico das áreas comuns do colégio, no Anexo A.
            \tA vistoria foi realizada em conjunto com os membros da equipe de manutenção do cliente, que visitaram os setores da edificação, com o objetivo de verificar os horários de funcionamentos das áreas, assim como os tipos de lâmpadas e reatores presentes. Com os dados da vistoria em mãos, foi possível então realizar o diagnóstico energético levando em consideração os cálculos preliminares do consumo Energético Anual, Demanda Retirada na Ponta e RCB para a dada proposta''',
        
        }
        return self.document_paragraphs[paragraph]

    def create_file(self):
        self.first_page()
        self.resumo_executivo()
        self.dados()

        print(self.name)
        # print(self.client)
        print('passou aqui')
        self.document.save(self.name+'.docx')
        convert(self.name+'.docx',self.name+'.pdf')
        os.remove(self.name+'.docx')
        
        webbrowser.open(self.name+'.pdf')        

    def first_page(self):
        self.document.add_picture(self.base64_images_decode(self.document_images['Logo']))
        self.last_paragrph = self.document.paragraphs[-1]
        self.last_paragrph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(21):
            if i == 8:
                paragraph = self.document.add_paragraph().add_run('Diagnóstico Energético')
                paragraph.font.size = Pt(30)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif i == 9:
                cliente = self.project_atributes['Cliente']
                paragraph = self.document.add_paragraph().add_run(f'Projeto de Eficientização do {cliente}')
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.font.size = Pt(15)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif i == 20:
                paragraph = self.document.add_paragraph(f'{self.month_name(self.date.month)} de {self.date.year}')
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:                           
                paragraph = self.document.add_paragraph('')
        # self.document.add_page_break()
    def resumo_executivo(self):
        self.document.add_heading('1. Resumo Executivo')
        # self.document.breaks()
        self.project_atributes
        print(self.project_atributes)
        print(self.set_paragraphs('Resumo Executivo'))
        paragraph = self.document.add_paragraph(self.set_paragraphs('Resumo Executivo'))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    def dados(self):
        self.document.add_heading('2. Dados',1)
        self.document.add_heading('2.1. Empresa Executora',2)
        #gerar tabela da empresa
        self.document.add_heading('2.2. Cliente Beneficiado',2)
        self.table_clients_write()
    #Function to make a client data table
    def table_clients_write(self):
        self.client_data = (
            ('Cliente',self.project_client['Nome Cliente']),
            ('Nome Fantasia',self.project_client['Nome Fantasia']),
            ('CNPJ',self.project_client['CNPJ']),
            ('N°s do Cliente',self.project_client['Num Cliente']),
            ('Modalidade da Tarifa',self.project_client['Modalidade da Tarifa']),
            ('Classe/Subclasse',self.project_client['Classe/Subclasse']),
            ('Endereço',self.project_client['Endereco']),
            ('Cidade',self.project_client['Cidade']),
            ('Estado',self.project_client['Estado']),
            ('Telefone',self.project_client['Telefone']),
            ('E-mail',self.project_client['E-mail']),
            ('Contato',self.project_client['Contato']),
            ('Ramo da Atividade',self.project_client['Ramo de Atividade'])
        )

        print(self.client_data)

        table = self.document.add_table(1,2)

        for data,info in self.client_data:

            print(data)
            print(info)
            row = table.add_row().cells
            row[0].text = data
            row[1].text = info
        print('Tabela de clientes criada')
        print(type(table))
        

#Creating a main application functions
class Application_Functions():
    def base64_images(self):
        self.ecosol_logo= "iVBORw0KGgoAAAANSUhEUgAAASwAAABkCAYAAAA8AQ3AAABX5klEQVR4nO19d5wV1dn/9zln5tbtlbK0BaSDVEWRBRXsNdw1GqNREzRF03x/Sd4k3r2kvOnmTVUTjYl9r8bE2KIiLIoNsCAsIEXpsLC7bLtt5pzn98edWRbchW0U8+7387ni3jtzzjNn5jzz9AfoQx/60IdPCOh4TcQAEcDXTcLAG0/1ftMrYdkaMAXlbKrXG699MvULIoD5eFHUhz704ZMG43hNVBEGIQLOCRgDZg4yvu7xEWADCBDkB9YaAL84XrT0oQ99+GRCHO8JPUSWpZGEDaUUUrBheyXVHW86+tCHPnzycNwkLBekIcCQAKQGIBlSp//uQx/60Icj4rhLWO2BCQwcR4NaH/rQh08kjjvDUvRxuzoBfab2PvShD0fFcWdYgsBEhzKoPm7Vhz70oTPoEsNiBjGffJobAyTSYtpJR1sf+tCH3kOnje5hQBBBA2nGdbiU1FnIbp53BBAB7HCrPmGtD334D0anJKzKEGQE0A9c7it/4ybfH4nAS8qOv4fxY3QBEgD/YZ45/uXP+h8rAfzMoD5Jqw99+M/EURlWOAxxVRTq/832jLngFPG704YZt7xyvXfh3CrYS8Injmkxg0IMnQdknTvK+OPcSUb5ozcE7iQCoxICfUyrD334j8MRGVYYEIsWQQ/LRfYtk+VD+dmyEDbscf3NPz51lfecuRHYHOpaDJXqHRsYIQpBBHruRv9dIwfIWXYzp84cKm5++hrf7VQOFQ73Maw+9OE/DR0yLAaoAuCCAmQ8EQr8bVixMdlKsLJsiNwgidMHGw/ef7l3OEWhwuHj623kSggqh3r2at/3ZgwzrlYJKFYwAaizhsofPvAp39WRCDQfZ7r60Ic+HFt0vKHDIAHwt6Z5BhRl4DwoaDBICAgrybowR/SbO1Q+snAqAhUV4PBxCpHgECSVQz1wuVl+1nAZgWatGUIQSNnQWdnCO7GQzgEAjOuTsvrQh/8kdMhkKAKtGeL251IfLNtiX9MQZzJNgBksBIRKsD24UE6/bUbgXiJwRQiEY2w3CoUgKQr1s3nmqeefYt6d4RVsWSBJIGYo6Sdz22715g+Xxb/GYQiUp72affgY6LDPyYaTnT7g5KbNxcm+hl3GEY3mRNAcgqRo6u8vZorwOSONRSRhQ8MAwVBJtscNEJ9+86ZANd0b+8GSMIy5EdjHglAGhHwcamYpihaMMx4uyJY5Vpy1FBBKQ5teyP0H1L573otdFa1GMyJOuEPvgsLh8DG9+ZFI5FgwWREOh93xGe2sCzNTRUUFHemYY4TWNa2oqGAiOtnoS0fMuDbRCrBIh+a0zq+59TcCAIrguNLXGRpBgNYgl0ZEwMciDIhDkBjrjFsNQhT6uIcbcWXasL7y8/5H+EcZbN+RYdl3ZLD1/Qzmigy76XsZ/PzVnk8BQEeeQ9fO9a0zzUnN/x1McUUGp76fYfOPMvmdm/3LAEB0wAoYIMce5Vm10Pcs/ziTU9/PsB0aNEcy7MbvZdiPXumbDTiL1gcAEKFQqL218ALIAzAQQCZw6HoRETo4r1cRCoUkM7d31/0AigD0B5CBwzSB40UfACwpg0Efp1AACAIIADAP/1FS+rxjTx2A9N5oby6BNH0BpNfz0JMAOOd9oqSvTi1qRTmYGTSA4jctuzVQOqK/nGHHWRFB2gqU4QFPGWz85S+X8ea5EevdyhBkeRSqs0QcbcWWlkHOjcB+6bPen0wZYlxgJ9gmgsEAIKEgYKzZpm759N8Ty7gybePq7NxdIJFnzZqV6/P5RjEzaa1JSqkdiQDuvx3BMIxDLlNrTbZtt56XTCYDSqm65cuXr0bP30gUDocpEonoaDSKESNGlBQWFk4zDOPSzMzMcV6vNwuAj5lNIkporeOxWGxvU1PTCx6P57lly5atiUajKhwOi2MkzYhwOIxIJKKICFOnTh1tGMZsv98/LzMzs9Tj8QSUUn4iIgAJ27ZjLS0tW1taWl6QUi5+7bXXNkSjUcXM5BzT61KpK7FQWmPwv3y9b8qBJJ83Ol/MGpBBQz2SvERgW4Obknr/5npsrGnRS4v98rVZf42tmVsFWxCg7oCgyLExTXAYQkSgKQI7C8h77kbf+Jpmnj+qQMwcmEVDPZJ8DDARkLQ5ub1Bf7ixFq/kB/Bi2f2J9ymCJiJA9xKNz94K73D2zahtYa+Q0FmZpnppjbXttheTH7oFPHs6R6e5q7s4v7/IW3rlGOPV4mzRP5Vidu1Hhpfk1v1q01/ejJ1Z8Rr2VVSAIm0WIRyGiESgv3WmOen7ZZ4VQQ+ZloIyfSTf3Wa/Ovnu+FmCAH3YJblq5mOXm1+4ZKLnHi+R0jod0qAZthkgY81W9bsJd8VuPUbMCmVlZUZVVZU9f/78a7Kysh5ymEyvvZm01uz1eunAgQMrn3322VlElHQkj+7cYAFnA5eUlIwfPXr0jV6vNxQMBkuEEOAOSrqSI0bEYrFUMpn8d11d3R9WrFjxPJCWhKLRaK+sa9uxJkyYcE7//v0/7/V6LwkGg0EAR6SPmRGPx1uSyeRzdXV19x4L+oCDzyoAPH+t7zMjCsTCoVlitvRSemUPp5Gc/0igoVFjT1y/uL2OH5r3QOKvQFpD6eXnkjjNTDWA3LcX+r6SGRALRuSIiTA7oBFIswwB6BRj8wG94kCMH5rxp/i9AJo5DNFdVdY5V9eFfYNzTXMr3FeIn7Ds/dQdZX9N/KC31qDTnj2KQD8WgvzSM8ktyzbb1x9o1rbHgNYMJoK0U6yGFMkR10zx/5UIogKdD97siPNyGGJuBPbPz5Nz5o4yf+s3iJUCEYFYQ5l+Mj7aa7844a7YbcyQdIyN7MwsHGaluReRHpo1Mxvogbc1HA67zCqjrKzsh1OmTHm1oKDg65mZmSVaa21ZlrZtW9u2rQ7/WJalLcvSXq/Xk5+ff8mgQYP+NX/+/D+fcsopBdFoVPWGChYOh0U0GlWDBw/uf+655/5l+PDhzxYUFHza7/cH3fk7oK/1N5/PF8zPz19QUlLyr/POO++hU045ZZgrDfaUPiD9zEUi0LdMRFH1rb7Hy4YbDw4vkrOJoKw4ayvF2ragLPeTgrZS0HaKtR1nle0jPaqfMW92qXH/h18LvPWzecaZVA7lmCl6/JILhyEEgSkCHQ2Z1+/+RuCdyYONRSMK5USloawEazvF2kpBW4fRaadYW3HWYKiRxXL69EHy19u+Hnj7/ks9CygCzXzQdNMdeEHaTrKtEqxsi5NIsgIh3tNrbosu6dnlUSgOw6BI8sUXP4NvzDnF/I0UcCUeqRNsn9LfOH/FTd5fUCT59SNw1UOrNTCbh38ZDkOIRdBfPc07NDRG/q0wU3pTCdaGa2T3kdxXqzf/7k1xLTNQcbiR8dhAEhE5TIYc2lVP53WYlXCki2491KFQSEYiETVu3LhJgwcPvjs7O/s0rTUsy1LumESkAUghhKA2hhlmhtbaZmahteZkMqmllEZ+fv5NpmmenpWV9dloNPpOTyQZl74pU6acNWDAgHszMzNHKqWQTCYVHCZNRJqIDCEO3TNaa2itFQAopVgppQ3DMPLy8q7xer1n5ObmfjkSiTzrqLDdfmm5ksIvLvCcEhplPDm4WI7VcdZWnJkACYJtChAEZOtdSj91tq3AGjBsDeIYK4NAQwvl9M9P8708vlB9mx6O3+nk4HZXckYYrZJfcNXN/t+f2l9eLySgUrCVZiEIEkirosKAcciTpAHNsJQGaYah4qwEAYPyxcirszzRkQXi10SJ/wJg91B9k5TWuhgEqbl3bXld5qYUgc1hGPMeSv72nW32XcIkCScpWgEGbLYnl5hfe/la701UDlV5uAFcQjsL20qBhtgLOF4M59uKCjAzfDdPE38dUiwHWQnHI8hg0wPUN+vkUxtToV++3lKDClDkGNgxjoDWR8Hj8Uifz2d4vd6efDw+n88QQgS6SY+MRqNqwoQJ55aWlj6bm5t7mmVZSinFcO6xYRjC4/EYWmuKx+NN8Xh8TywW2xWPx/clEgltmqbh8XhcRiaYmZPJpMrKyho3YMCAf40bN+5Mh1l1+ZlxGd3UqVMXDBo06JmMjIyRqVRKaa0ZABERPB6PME3TsCwLsVjsgEPbzlgsVptKpWCapvR4PNKhj5iZU6mU8vv9Q/v37//Ps84666ZIJKK7Kwkyg1ABDo1F3oJRRnRwsRxrx9jWaV8QGRIwM8hgQNY0cXx/k96/r0nX1TSxDcAwMsj0eIjA0ASQBoQVZ53rh3HBGONXb9zk/6Wj1QLdeCmFwxCLCPrWGcj64Db/P6YMM65nnZacmCEFAYYEzCAZwoBR08RqX5OubaWzmVmYMM0gGYYAE0DMEFYS2gDUGSPNr635kv/BIYDPoN6Lq+xukYSO0D3uVwHFFRBEia9Vf5lGjRkg51oJKJE2wkvDID2lRP7mwSt968ujieWVIchyxyawz7Y+akh4dvu8VCI0a21DKq2edcaVABSH0xe6/Ebf78YMNGbbMVZCQDKDTQmV0jCWb0197vNPWe+kJb5jE0pxOIQQbZkiA6Da2tp3tNZbhBBCaw0cvNEfq1PYnkfMkXrYMAxPPB5/A0CX7FcuM5gyZcpZJSUljwYCgfxkMmkTkcHMSkophRDU2Ni407Ksv6dSqbc2b978QWNjY10ymeRgMOgZPHjwsKKiomlCiCuzsrImERFs21ZCCJlMJlUgEBhYWlr6mJTy/NWrV6/piiTj0jd9+vTzBg4ceL/X6w1alqWISDKzMgxDMjPq6+vf11r/o6Gh4d0tW7ZsaWxsbPJ6veTz+QIDBw4cWlRUNA7AxZmZmWeYpknuGLZta4/HYxQWFv7xzDPPbIpGo5XdkgQr0jahlQv9vx5SJCfacbaIYDKgDB9kbQOr/fvVg7UxvPXAu9b7r+2x9+YK0zNzqCi5bKQ4pTiLhiqNK0cUiyFgQFlQJCAtC2wKVqcNld948VpvM1GygtMiTKfvsVvSKRJBxm2n+Z8e0V+epWJsIe2hZClA5AHtO6BVzV6OCsJbj6xR61/4IPVRXED5tMe4fAxGLBgnRluM0/oH5ZW52UQqCQaDNENwnO1xg4yrnlwY8E+5J1ZeEYZVUdH9qiyH095b6PZg7KhsXxyDft86L/D64Hwx1ErgYFyUj8SuWrXjv/8dm37/WuytIFBFJYjKod79YuB3kwbLLyPJ3BjXDQufjk18bC22t/HM6Geu9nz1/DGeX2sFBQ3ptACzDT8ZK7akfjDjz8k7jhezco3u55133mezs7P/ppRSQghpWRZWr1594ZYtW57DQWN3W2WhLTpaa/c4CXTZKCkA6PHjx48eNmzYkmAw2C+VSmkhhGBm7fF4RCwWq2toaPhlS0vLfStWrNhzpMFyc3Ozx44duyAvL+/7GRkZQyzL0gCE1lp5vV5ZX1//9vPPP38OgIYOrrFd+iZPnjxm0KBBy/x+f4FlWS59yjRN2dzcvHn//v2L1q1b91RDQ8OBo4wXnDZt2nnFxcXfz8zMPNW2bQ1H2pJSCtu2G7du3Tpv1apVb3WNqUJGo1D3XeY7a8E4uSzoIa3TvQfY8BJtrrHf+8d660u3v2C/dqRxbjsNxWcP8V44daD8Xkm+KNU2tFIgBtjjAdc3a/uf72Pojc/F9txxx0HD/tHgmlbW3uJ/cOwQ4zN2zPGSa2jDA9GSYr1xv/3gup36Z9f8y1p7tPHuvticclqJ/P6YfsblHgIsG+wwUNvwkvHmZus3p9+X+GpnDeWuKt0S9pd4tNwmHIeY4Sdj6QbrO3P/mvjJcTe6Hw6KQD+2APIP1dizeGOqvL5ZN5tegnIj4ZNsDyiQJV883XcjEXhOGCIaTZ+7qVb9K9HCFnuIdjfxisfWYjszqAIgsQj6f8/3l0wdYEQEEbOCcBRqywjA2LDDenLGn5N3cCUkRXrfI3gkaK35MC8WB4PBFiJCZWUlERGIiJ0PDvtwBx/3d9XWrtQJUDgcRmFhYcagQYPuy8jI6JdKpZRjn1KmaYqGhobXN27ceM7SpUt/vGLFij3hcFiUlZW5hv3WTygUkuFw2Kivr29Yvnz5vW+99dbc2traZw3DEETkMmc7Nzd3yty5c78NgDsRQEuhUIgABIuKiv4UDAYLLMtSLjM1DEPW1tY+8dprr8154403/tbQ0HCgsrJSdkSf833LypUr//7MM8/M2bdv313CMXYRkVBKKY/Hk9W/f/8/lZaWZldUVHAHMV4fQ2VlmvGeMUjcmpkpoNNpaMrwEm2pUSs/Vxmfd/sL9mvMECsXwgyFIMOACAOiMgS5ciFMZsjfvIm9l1cm/zLoztiUt7fZf9DMwhDpFDeA5Hu77e/d8Gxsj34MsqvM6tXrvbeMHSg/o+NpZqU1tOEjsb9F17z8gfWpyXclr7/mX9ZaZtDKhTAr29B4GJ1089PW26felbji39XWF+riHDc9RMzQYBjaYjVpkHHbI1eaV7VxFnQbupclrB4ZxMqj6QuiqL2iX7bxxXNKxQMeA7al0gFpEuCB2eKCEPDTORVQ7n5cUJlcXHO73OrLlCN2NfFDcKovVISASBRUmmuXFed6s+1kWhXUDGX6ydxVa793+5OJa5lBFRXHP5q4nVgrZmZiZvr973/frsrXRXT6ekKhkIhEImr27NnfzM7OnmlZli2EMFxm0NDQ8MzKlSs/vW/fvmZHRdKOxPGxjRJ13yRpJiOi0eiHL7zwwuXz589/MC8vr9w13CuldE5Ozu2zZ8++NxKJbDySFOOMo2bNmvXl3NzcMx1mKl3Jat++ffcvXrz48wCUS195eXm7L6DD6Xv88ccbFi9e/MW5c+fWFBQU3IG0ZCpTqZTKzMycWFpaegcRfdOxZx3xpRYOpwtTPvxp/4A8H87nJKezKiToQIuOP/S+9flX92DfkjAMItgfW780aQr3pNWfaDnE1Y+jYeo9iS+/fL1/wxkl8k5vkOSGHer1uX9L/doxvHeKWYXDEFgL/tosb+mIIvk/ALRiSDC06YPYXa/33P1W6rLIK/ZbS8IwlgLaUeGsjy9iK51uELagSPLP0cuxbfZI48miTArYFlgpkM9LPLfU/J9ZI6yXEcX+3oqh6g302LBGjufwwocSD67aaf0IggwTSIFB2mJk+2jGFZ8yRxOBw+G0oY8Ae39CP5Ns1PHV9YklROCKta0Lwqfky3IQMafjsizTQ7KhSe3/9zoVemYPYjgsxutEwk0lKSoqchloTz6dghseMHHixFPy8/O/qtOin2Rm7UhWb7z22mvXtmFWnfVksnOsJCLrhRdeuKm+vn65aZpSSimVUqKpqeme+vr6/QDICSptl77Kyko9derUwXl5ebc7Yqmrpsra2tqnFy9efDMzazgOg67Q53hUxZIlS8L19fV3SiklM2shBCmldGZm5i0TJ04c35lwhwrnXxVTkwszRabWYBBYmETbD+gn71hivceV6cDloxFHBC6PQilOR5+f/df4b97erb+994BOvb3X+gIBdrQcopPXioqKtCZz43iEi3NljpUCEwDpITTFdPPja1R55BX7rZULYc6NwHb2xNHGZnKCTdeE4Qn9I/nC4k3W1Q1xtoVBmgjCTrIqzpfD/vds//9zqvmeNNHwveMJiKSZ1hn3Jb+/eqf1JDLIY5ogpWBnBoU5NEdeSgAqqkEoBzGAHc20fHWNqv7a09im74BABBBRqOuno19hQM7QilkKsBkgsyGmUy9vsa+78ZnURv0Y5LGKHP6kobi4+Ls+ny/X8QbCMAyKx+N7duzY8ZmGhoYDPQhDUMwsATRv27bthlgsVpdIJPZu27bt2pdeeulL77//fr1zXIebg4g4Kyvry8FgsNC2bZc+0dzcvH39+vU3E1HKiVLvDn3aCSQVL7300n83NTWt8Hg8gpmhlGKfzxfo37//tzoz0FJnD5Tk0lkQYABaMBhgNCbxd+6enZedPSHOuC/2i/tXJs+65glrLQPU2QyQyhCkIOg/X2pOHpRjXIsUa8BhdszipY3qO7f9O/HKkjCMafe0I1F1AuMjSK1cCPOaJ1NPrd1t/1RIlmBoBqROMQ/NoZu+VYahRNC95TXsKXqLCEYFFAF86SOJG5evt363p1HHTR+ZkKDhOeISBgQqwRRNM5vwsnjVb1dY32aAKAKuqEyXNr5hnP+s3CzRTwiIlMXGpp3Wk39dlTj/ysdSz1WGjk0kew/QG2pgl+CqYRMmTJgQDAavdJiBICLWWqO+vv6b77///paysjKjh9HfKhwOi3Xr1m3ct29faN26dXNXrFjxkHO9R/JwUSQS0aWlpUWZmZnXOaELcOjj+vr6b23fvn3XggULJHoWiqIrKioAILF3797/SiaTKdekZds2BwKBy0499dSxkUhEH0nKmuP8OyRHDIIkYgZLCZGIAQHJuwnggxppl5CWZAD+9sv2W463rPMqf2U6Vn1yP/nFnCwhLJWWwg0fiS379atXRpN3cyXknB7acafeDZsZ4tdvJRftrFVrDC8RAVrb0Hk5MnfBSP8NAFBReXJIWb3GNYnAGqCtDTgw677ErS9sticv22z9prlR1xVkiJlvfcE/lQi6MpSe8/UtqHngXfUSuepQWiXkARm4OcaMdbvUS0+sS5aN/F3iyq++oJZwGKIr+YnHC8ebYbkoLi6+1u/3Z+h0LIU2TVM0NTUtWbZs2cPhcFhUVVX1eK0c+xS9/vrrL2/cuHFdOBx2I/073HihUCgtsZSUXO73+/s5DJVN0xQtLS2vvfrqq48BoGg02mMp2Y27WrlyZVUsFnvMMAwBgLXW2ufzZebn55cDQHV19VHvkVeSBwcvTMRtxp4YdUtyaQs3cb+L4QFEBD2zGEVDcsQVsBkEkBSgeFxj+Vb9XQBWNNrzxitEYJSDotVILdtm/4AVk6C08MCKMTCbPgPA7wgKXX7WqZeTq3tVzHMWj7gS8vp/pD4ouy/x1YsfaDmtukY9snqvuhAAQmNbF5jcNIAw0m7Ri0b6B9bGOf70u8mrxv4hPu/av9vLmEGVoZNaDTyuDMthIn6fz3elG3gphCDLsnR9ff0vgNYN2ltGUnaM150KE6isrNQA4Pf7rxBCMAAmIrJtG42Njb8HoB2m1iv0jR07lgHQvn37fp9KpZKu51BrzYFA4EoAnZI0lU6/DAmA1lDZAaBfQGdzGCLUMxK5q8/uknDaM/eri70X5WSIAq2gGGDhJdp1QFd97qnEq735AqcoFDPomiesJ7fVqTXCk96XbIELAmL4I5/ynAMArrBxIuF6CYl7owa6U2OHytP6OwBQBJsm/DF+zRUTUOL87d48jqSTLRFxHt4cjjfNvBdXArBaz0+/mZQbo9VTEisA9KbB3jCOXx+ONhHjszweT6nWmgBoKaVsbGxc+eabb74AAL2ZCNzF8YiIeNCgQQMCgcBpSikCACmliMViW1999dVnnfF6bf1dJrpq1aq3iouL38jOzi5LpVLQWpNhGKMnTZp02nvvvbf8aPa8uqRuGKQlyEnAN0yBmBLnUwQvdVC+5ZhhTgU0IkCmT8yVHoIdYy0IUqcYWw/ovxKgly6Fgd7M7qiAJMA6kOK/DBb0SwHWigEzQMaUAcYFQOrpUAiut/GEwb0R7GRq9xgchoEIlMuYwoCoSPtFdxzpNAB4aBMaBQGPLjgkxoqWuAGivURjd9GO+ndC1MGioqLZXq9XpFIp26UhmUw+jIMhAidEdS4rK5NVVVV2SUlJmZQyl5mZiFhKiVgs9gKAhp7m+3Uwr1FVVWXbtv0PZi5zvlYej8fs37//jPfee295TU1N+/eqOv1Mrd/HKyb154VgCCaQToEn9afP//Jc868Usd7nu2HSzbBx7N37RAS9cCoCxQFxLpIAM6RhQDQndMNdbySfZgBzq3rXPFLhFEp5YJV+/vtZ+gfZAQooGwoakIJnhcfCI65CCke2Xx5zGABw9VQUTMo2z/V5ANZgzSBBrVULPxb8JQjcdrWIQYakll27U29SBPsIaS9HKF1tUEccnbgT+jZpTsd3wZX6KsBO/Iv3zxd7ZzVZuoiIDnExa2aSh9Ep6OB1AIAFSA14tUV7v/1y6jn00sIfRxsWPf744wrpyPMz3S+llDKRSMT37du3FDgkZum4wwntQG5u7gTTNN3Ea+HU/XoJAHXGntRVVFVVaQDYv3//G8FgMGmaplcpZZumCWY+DQAtXbq0/cBcx0SRYP16LMnwGSQEA1qxzvCL7KsmmX/fHUOIbrbedV+mznMNHIONGw6DIhFwUdAcFTRR5JSJYRiEbQ38ZnQT9vWkkXFHiKSrNRBRqvobs+TmbGFMEMTgFGNAUIwdPtw7kKuTH7r09ebcXYEBAGcN8E787ATjkQw/Aa3v7C5CAtvqjB0Xnqp/d9b9iT+XR1ELoLVLdCeHbLsQaakvArx4redTg/LkV0flybO6QRkcOkB+QrxJrfj2y6nn3M47XRkjHUR+6JUcrXBfb4KZMXr06Fyv1ztOKeVMTwCwafXq1Wucw06YY8JhqMTM4936VVJKkUgkGjdu3Lga6TiqY7FeGgDefvvt94YMGbKViE5hZlJKwTCMqSUlJT4iiqO9l1QkXZySyNo06zbzzeH9xAwrnlbBrCT0wFwx4vYzzJfPGyb/Z95DibvLo2gE0hnR6g4Y5dXgaC86gyqqQREAZUPETI+HpNLplz4kpNL0KgBycm57PyWtIi1UbKzTiwcUYAKnAAbYZ5DRxDwJwIcVACK9PnHnYQCACUombXCGAux04FyXWRYrYHCeKBmcL35Sc3vg86/vsv/82Pupe4hQjy5KM07eoPHb+d6580eK7w7KFWX+gIAdZ90d2hwoQ0EKgWQ3z28XQgi3DLEMhXpmnnU2c0fqEgFgv98/0DCMHGeXsRACLS0t1QCsY6FudQHEzFxYWBgkolFOIjiEEBBCfLR169bNDhM7NtU309VH48lkcqvX6z0lnQHE8Hg8+cXFxTk7duxoty4TAczlkACSa/ZYvy3NFw8KkX7BCYKwEtDFWSK3OFP8bP9/BW98c6f9xAe1/MDX/53a5OaxMkOkjaO9UCd9bPr5nlQsS6WPYLewcpq+IKH1BwB46dIezXAkCAb0gAxa5bj3mAEtTMjJxWIigH8sbVMg8kTAAADNMEQbRtAdVyQRYKXAArAKs8SISwu8P5nSzwiPLbJu+d5i62+dLZvsSGTev1/l+dcVUzznQgFIsZWKMUnRfeOnEwBIzL3r6WhubrYcm9ExlWxCoRBFo1EIIU4RQnjcpEYiglLqbaBz7vtjjZEjR2Z5PJ5irXVrFVPLsvYAsLTWdKwkUqdBBcfj8XXZ2dnzmNkVnTJs2x4EYDc6eHFSFDochrg8Yj32/heN68cPNubZMbZAMIWAsCywJOj8bDH6wizzu2c283fPGiqf90k888Nl9jIia7U7lpue020P3rg0fU1JLih0iBUEub+FrcfX2R8CwL6iY6OSLXX+rY3z3sEtDDNdzklLgzA6nwYBB+PWThQMACAJrRnaievn7pSEcHV66YEHDOzdr2r2NOE3pxZjtTOe7syoROCxY6EHZ4pFb2yyNo0rEJ/JDIpMjwAsC4qA7j3xBK3RWry1R3BKDbNhGDRs2LCfDBkyZIMQglyp4pDLSRf8c2uPt10BQjoXkYmItNa+xsbGv6xYseLp9iQl12js8/kGOvNrrTVprREIBD7o6TX1FE4NeTZNMyCl9LUWfiJCIpHYBRxkKsdifpdZE9EW5ytiZhZCmIZhDAAOMv12kA4PIdj3vyuu/bLHXjysnzHejrPt1JoiDUiVZE0Asnwkpg41zofi8+88n+p+Md98f/GW1KPa9PyDKLYHgOJKSJR3uWMMyaugAIiEokHpzQiCAKTg+qfW2dsBoDx6bCScpc7eeLxa7x6ao+P9soVfWWnGazPlAgAcD+aJggEAgkEZHhLCQxDdCg9D+rYYwM5atXdPI//l8kfjP9/RiDoAwMNdG6q6Gqlp1YlXALxyy0Tz91+ZJb9WGBBXF+XLQLdtbAwBD8FIwNeNsw8d6mC0N3Jzc2cCmNnD8WAYBmzb3gjg6aVLl3YodgcCgTxHqgKQrsYJINaT+XsTLS0tufn5+Yd0krFtey9wfCTArKys3W3+ZCIin89XfLTzIpG0lBWJtNRsa8BF4bn++8b1l+dAAcpOR4MLkZbOFQMcYyUI3C9b5IGo7LocX9n+Jv2LFTf7Hlq/x36Qyu1XAKQb/3ZB2nK4mwEg21ULQIApKL5+Pw50eiG6A4cRPfOeVfv/ZnnqQBjo/tScQA4AmLK1hNIJMbwbAPBBrapd8hGW5HnBFgPpXKqOodMlUMlVI4mhNLE3xXhl8Qbc+6PXkpsBgMMQFehe3FNlCDI0FkwRa81dq63PP3yJ+ZviPH2zT/IEYkopfDx6uK1kmA7iAjHYoZVSQpAnYfNGAMDa3llwpyZTT8dSRCSJjh5Z7fP5AkIIt1yNYVkWGhoaEj2cvzeRTUTCUckIALTWB471pG3CFpJtpDu3fI/3sGPaRSQC7dR22hatjl/00nW+r00qFt8pyBbZEgQ7yTannztyq+baKWinILAoCIpgQb5cWJonr68uMf7129cSd1AU6xymddTnJB11DQAwDYHMQ6qrpSswJB377rFRqwGOAFgfQwqA3cqXGFAEDwDB3LW9fEzKy/xsubX2Z8uts3trUKewXmss1kHmc/SL5UrIivJ01jvgxHGFIShirQasL/cWjT2JnD9Y1h2unaanN+VwdbFDaK0PscG5Lq4ezt9r6CDM44SmVB1eI/5IoAg0AyQJyXP/lvjpeaPw2M/m+L6eGxSXDsoXQ8GUrnquoW0nUVgAAgTYGuBmtvO85M0bYCyIzAuef8kY+5v0cPIehxECnXi5zciDFARf2yOJoJAWqHs9pKF1DrQ6vFICsNo+kax7p4lGT9F6J4l6/uEwhHNj3AA7YsfY3kkGQVQOFQG0Wws+gnQpjMoQpGOQ7xmdvb2A6d0giKjbHwAmEQk3Ufgo8x1SppmIIKX8WDPPE4W2RnX3/4nouNEnxMF2vG0kvC6NQQA7JWLEvzfgo0l3J7566z9js19cb1/3znZ7ccLiZhsQRoAM0wQxw9acDt0RAoalwFYCqjAoMuaPNu+uus7/U0pLb516/N6qg1aMRFvFi9LOota4wmMBRtotmAl4NGAeMhN1qnTNMUer162DdnBdQmt0exiiogIggqYo1P/MNif7fJjw9ResBxhom7Dp3hICwFNzkfWLS7xXzv1b8p/l0bT9q9V4GYU68fz9EDAAqqmpeVRrvRKuA+Og+HXI7XYY08GT07Z2dt38WmtvLBZ7BgDmzJmjq6qq2p00Ho83BQIBdww2TZMyMjK627yi12EYRh3S9re215t7rOd1g1Yty/J7PJ7W7x1JuMsqMwHsZFYQhyEpgu3/3Jx4AMADZ5Zi8O3TPFdOKzHOBmN2SZHIhk1pkUtDEUESIC2b2ZTQs0ca/2/Z53wxiiQibjnhDuZ0N0TKVtzc9nnXgAeAn9P2ymNqQ8oLwGSGcZBbAoJgwXWcnUC21es5Uk7DCRWJAJ8Zjf5fOdP/9dJ+4uY1W9XPAYuRZkCuiuBeOjODygnNhUHxjdrvBm5/ZbP958erU/dQOVqcNTppqh5qreEkHGP79u2/3bx58xFrfXcVR4qlisfjdc4mFE4tc2itM3pz/u4gEml1HTUyc0oI0eopNAyjCGhNVD6miMVi/TIyMlzpTmit0dTUVAMcZGpdBFMEdjgMUTEOJK6CWr4F25ZvSf0aSP36J+eaE6b0E2VDssWNg/PkZJ+PpE6y0o53MaUgPMRqYj/jjt9fbCyliF11pBAfQYBiWCA0tL7ZGbA0+yfnI+edWsTCAEWOwV6IhiAQhVowySwAc27bcO9MM23wtx89sSWeepVhOW8PddYADPr5Bb4bBueIW/sXyIKGBt24ck/iPgLQprIosrKQ19CIegIYUYgooL6t9K/ycj33XTZG/OrUYuNLN05UPz/7weSjBDS27ch7siAzMzNQVlZm+P1+GY/He3QjnRSTdq/P3WzxeHynSneTJUr38UM8Hh/Vk3l7E0KImNY6YRiGT2tNzAy/3z/w6Gf2DC4z1FqPcKPsnQqkSmu9p+0x3UEkAt3Gm09LyiDnLIUmst4H8D6AP/zhInnxBSO8tw/NF2eRBpQCSwJZNpCdRWJ+qRmpDNnnhCo7DPFhO93jU/sNsR0iHeYDBrSmnLNPMQa+87q9C2Ec03DzBWNF/+wABWCj1WBqSq4FAKw9sXpOrzEspwSMin7KLJ893PPTohwxFEkGNFRjUr/xrZewyymozwDwvdmeMaf2w9epMrUwHIaocL5/Zn3y1WE5sj43Q2QNyRcjhhTJu1d/SVz/3y/FL1u0CPtPMqbFlmXZVVVVdigU4ueff/6YvXnclBbbtjdqrVOGYXjcOCzDMKYA6dIuXWxk0ZtgANiwYUNTTk7OPq/Xm4O01gqPx9MfgCcSiRyz5Fm3XLPf7x/t2KzI8RA2Wpa1te0xvQCeWwUblC59FKoECwH9pWfUU0DsucXX+b41c5D8gd8kshVYpJsM86AsOfvpdcY0IvvNDlXDmyEB6Fy/rgEEiACloHOC5Ln2VHPoL1+3V8wBROQYRJu71RiKAlTkCwjYMdZEELAZG2v1TuBgcGln0ds2t16J+g6XwSiPQkVmGafNHWneW5QrhtpxtmwNBWa5tkZHGSBE01IYAPr0eDqtbJBRHi5DzqJF6bcXhyHuWIaN9Ql+F5JEyoKtkpyaMNg44wdzfH9lhlkxLp1I3Rt0f8LAAGCa5nalVK3LmBwJZgIA3/HMa2wPRITa2tomrXW1lLKVPqXU0FGjRrlS4LG4dy4TzPJ6vcNcVVQIgUQisde27bpjMCeAtDebKN3m3XEM2ef8LfHDlz+0P58EM4m0sVoD2usHnTfSXOCc2+7eW9o/fZ9X79IfqgSDAMmAkj6CR/ApwLGNNicAu5p4ahs2I7QFrNol3geAfdUn1izTY4YVBsSiZbAvGomBN87wPJSfKTNSLWwDMIWE2N+gG7fXYDEBXDE2XUoZABOLcwuKZPaMAZ45zCCuBKEaRAA21OlHAafyIcOjYmyfOti4sOo6//9QORTCPWs99EkFEWHVqlVNqVRqrVOHi52whmFTp06dBKRrZp0g8nj27Nmu42GdmzmklNI+ny8waNCgCUBr669ehVvldPr06acy8yBHwmIhBJRSKzdt2pR0yiQf083mMC/wEhgXP5y8d+Nufb/0kWROMzQmQp5fTAAAZx98bC2WOpLT4h36zVSKlRQQzBBQAGs6CwDEomPUi7M8XWKmNE/MgwYYEIJAScVsGvo9AFg79pPNsKgiBGKG76fz/A+UFMnhVpyVFDDAsIWX0JzilQtfTH7IlZCRCJgIPH0A8vtl0HwQeHCOvAZIl0iucBZjaxNe2FvPzYaRbkLJDAkbesYQ+c0nQ+ZNFIHNlcefabXTl/B4gt066ET0msOooJRSfr/fKCgoOOdEEXY46uvr37EsC45nVDvMdT4AdiuSHgvk5ubO9Pv9htbaLcMDACsBwMkeOB5g/CHtRHpjh/37pkadFGmmQ6SBZouHDclGjqD2HfMVTumWUm2tjVl6v5PYS7AZg7No5rUTUcTcmhvbi0RDEMDfm40xWV4xDCrt6SaTsLuRq59Yl9xObeg7UejJTSRmEEWhXrne86txJcZcJ7NcMgAhQZxi+rCO7yEAiB4s/frLeb4LMwKiAAmmkiwqK+uPAopARyKAZogv/jP50f4Y/xue9ANPaW8L+TzQs4d5fvOnS4wz21Y1PV5oLyjSMIzjpp66eXB1dXWvJpNJJiID6Zw5mKZ5NQCvk4h9QlRmt478rl27XrEsq86JiSLbthEIBM7Nzs7OddTWXqXPuWbDMIwFbdRBkUql4nv27Hndoe242T3dRitfeMZanbJ5g/BQuhehZvgMFH9uui+LAVS0E5dFjsf8hiokdjbSEngpbXNX0H6/yPrcZO+FAIDKXn72K9PazcWjfJcEM0TQtp24q7SV+5XnNyGpKyG76qnv7Uj3bl80V6a9GU+We2+dVer5ok6ygpOuQAxb+Ei+vdP60dkPJB7TYQiKQs9xxN28AF3o8REpC6nMgCj60QW+8wCAyyAqKtIpNv9aq27fvldtM70ktYaWBLJTQF6WCJw33PPQrTNQIhcd3/ZDUkqcQKM24Kzfm2+++aplWRullK4hnv1+//jTTjvtIuCgitRbCIfD4mj9/VwwM23btm13IpFY5tLnNIUYOGnSpCuOAX0SAM2cObPM7/dPdluKSSnZsqzq9957b6VL2lFpD0OEe+clyIhCEGD5TDS0+RYGQQ6VfMQ5lqbrXSGp1YtWiiEIpBhseAhDs8V1AARCvSzphNKVoYv84jpCmtGYBKHijLdr7BcBYOnvT7ztuFs3x223dfdF3nPnjTB+pglaaQgCwBq2DJKxZZd6cto9rS3lmZEOLL3hVBT2y6TzkGJohpBeQsCgSwEAS6EiEWhUgL5TlfzohY2pT9c367hpgjSD3SaPgwrlkJun+B/SDF9FuPfF445wEpRIdhtCxGOxWNSpAMHMrKWUKCgo+C8AXrcxQy/NSZFIRLsdao5GX3l5uQCARCLxhFLKrVbBQghkZWV9sZfpo3A4zAA4Kyvrqx6PR7qBuERELS0tf0e66YVEJxiWI+X3qiR2yKQEaEDvTx3ZOTLXoeEbS5L/amrSdVJCEkAqwTwgS86560JjltOBqlfMIhyCJAI/dY35qSF5YpxKsiYCYJLY18zbQo+l/g30flnm7qDLDIvDEFdFoW6ajBGXjJF/DQaET6UASSCtoYwAjF377TXTfhu7gRmM8rRo6Yq/+QFzYMCgTCe+naAAIgwCAJluZJnO56qE/PzT9utvfKS+mFQgKaE4zbSklWQ1rkTOfv0G328pAu2Ixyec+x9P7Nmz56F4PN4ipRQAhGVZOhgMnn7GGWfcGIlEdFlZWY8fZtdQPWPGjPkTJ06cHI1G3TLDHa6122Bi48aNz8Xj8Y+klASAHPqmnXHGGTc4zK/HkkwoFBKRSERPnz79vGAweJFlWQyAnKYXjXV1dVHg6PFX4XRSL/32As8pt85ACZHTJr5HxEEzICybAoCTSUKEuIW6yg2JZgBHiqXSHIJ8dRP27WziJ+AhIG24134/0fwRnh8CkKFQr7ysCZVgAOakQuO/hUFgdgoRSmB7g34YQItjMz7hgdtduimMdI11BrL+a5bvof75YoCVTHsyNEObHhL1zbr2wVXxBQcIDdHytCGv7RgyXQ/+kO8EfbwKKJWnawpd+Gjir+/tVD8XBhlEUOlcCUhlQU0dYnz+6U97vk7lUEuOg+fwRPUgbAu3/fratWvXNTc3P+YwBO3EHHFRUVFk/Pjxo53YsG6viVuTa/z48aP79etXOWzYsBfLyspuZbfGeMcbhcPhsNi1a1dtS0vL/W5un8PouKCgoGLcuHGDOtNG/igQlZWVevDgwbnFxcV3mqYp3JQoKSXF4/Hou+++u7ETVVipIpROKD57mPHXhdP8DzJDVFSAu2tuaG1fV+YttTSPhO087gbgEdi6YhcOMB85Wj2K9H57dZf9x8ZmHScDAoBQCdZDC+RZL33G/xUqh1q1sGexlLwQBhH0W5/3hgcXGxNVkhUDwjBBjY266akP1H0A4MZJdnn87lcIbhdduSGEtN2K3rrR+8dRA4wZqRiUIEjNYDKgk7bWyz+0r/vWEmx4bEHH6QeHX0FHVUBdpnXavfH/t26X+pf0kQFOu461hjAF9Myh5s//cqnnorkR2NxLInIXcOJchsy0Z8+en8disSZHyoJt2/D5fIVDhw59YMCAAfnRaFSVlZV1+YEuKyszFi1apIuLi4ODBg26LxAIZHu93vyioqLfzJs379Hzzz+/0A1baO98J02HPvjggz83NzfvMQyDmBm2bXMwGCweNGjQ/QC8ixYt6oya+TGEQiEZCoWIiGjUqFF3ZWRkjLFtWxERSSlFPB5vqa+vv5OZj9r0gkMQFIV69jO+68cUi9PH95dlqxb67yMC/UCkJf2u0lfhlD46bzhfUZQns5QNW6YjxVhK3ghAOXXZO3x+yqNQCIO+9JT1zs46/XfpJQGn3R2B9eQS8eM7L/LOmXYPrO62IeMlMOgeWP+8ynPx2H7md5FixWmboIZBYm2N/ZcfLktt5HD3A1V728Dc6fE4lG7a8OzVvu9MLzWvseNpZgUAJGBLk4yVW63vXPJI6lkOpwNJe4PAivK01+RPb8c+t6NGVRtekprTRngrBeQFhTx3uHH/j872jKUoVC8ZTTuCKy2cUNE4EonoiooKWrNmzfqGhoYfCiHIrVxqWZbOyMiYNnHixCcHDBgwqKqqynaYVmfedOS2zGLmzIkTJz6YnZ0907Is5TAElZWVddWBAweGERGHw+GOxtShUEjs2LFjZ1NT0w+cuZmIRCqV0jk5OWfPmzfvL8zsb8NUO/UmLisrM6LRqIpGo+qcc875VV5eXrlKG8skM2shBDU0NPx6xYoVaysqKuhILc+YQaiEfuAKlEzrL35OBLZTsKcMNq5//0v+B7RGhvvS7Ky05WR82GcNRv8hOcbXYDNz+lyJFNMLW9S/ASDamQDMSJrGx6r19+ob1D7DQ4IBWBYoL0MEPjPWeOS+i43pFIHNYRidpTEMCK6EpLmwH7nMe87ZI4y/BT2kbQWhATY9JOoP6N2b6pM/YgZVHMM0oK6iU5zZrZp43yWeK2cPlz+EgmaGEARohjIDZK7fZt8766+pnzvpNx0+JKqLbs4IoCsqIO58A3UjcqyrF0zAK0WZIjOVAksBYSVZlxTKgtBYfqi6GWUVFWiqiBybROmPtcxJhxScEDUxEomwo+787/z58+cXFBSck0gklBBCWpalcnNzz5oyZcpz+fn5X6qqqloGpKWyiooKqq6uJteu40ogY8eO5UgkoquqquyRI0eWlpaW/ik3N/dsy7Lc4oK2lNLYv3//nW+88cZbR1O1otGodo65e/78+Rfm5+dflEwmlRBC2rZt5+fnX33++efnbt269ZaqqqqtLn2u0f5I9A0ZMqTf6NGjf5qTk3Od1loxs9RaK4/HIxsbG99ftWrVr5y5j7yI0fRLuOpz/p8X5slCO842EQzbYjV+oHHNpq/5Ri770P4alduvpelzCiqUg6JIdzGPVoNCABACEIImgsrNRfbdF/v+MiBfDrASrAGQYQI1jbyzvi71EhFwhHzCVhCguRwyUpX8aEp/z48vGSvuFOmWd4aVZC7Mpn4XjfE89zcpbqZI6olWGqMQiAIYC65AunkwqkFtaNSRcuDJq7w3zRoif5vhE34ryUwCJAXslM3GSx9at133JGo+W9GzNCBmpxhZL3k1j8qwQg6z+tHZ5sTLxhp/CnoFW0mGFCCVZlZy+1712pg/xG9jhjgWdXMoAr0kDGNuxFrdL4c+N3+4+fcMD2nbBpyOIvbI/sapX53ku5coEeIQJDpR4bE7pBzyBxFSqZQfOJjrdxzBkUiEiCi5fv36L0ycOPGlzMzM0lQqpYhIplIpnZGRMa60tPT5goKCe9asWfM7ItqEI69JsKys7Mbs7OxvBYPBge5YAGzTNI36+vrl7733XpiIOpOXx5FIhACo9evX3zphwoSxmZmZwyzL0k6lVJ2bm3u+aZqvFBYW/nLZsmX3EVETjlzszztz5sxr8vPzv52RkXGKU+1VMrM2TVMmEonGXbt23dTY2FgXiUSO3t3FScSvj9lP1zeK8lw/ScsGC0DaCVbDC83pOV65ZGWx5y93vmr/hihV3eH6OaXif32+MeuKUeb/Di6UU+wEsyAIZihIkhtr7Xu+XoUDXAnpFOQ7KiiaVkupPPW7txeKWZOHGZ9ScVhCwLQS0EWZMv/K8eLxccXy/t+/rn5ClNpw+Bq2sm2Hxq9ON0/9wjTjO6OKZLkhACsJ7ZR/tqSXzDc3Wr8rj6Ye59AhDY27BduxYUXLP9bPoC2O9CxR22OOyOPDgFhE0GcOQu6DVwaWDSmS4+04u3YrZXgg65p553deajnjT29j2x3cMTd2k5ZvP8M8tWKO562gh0xLQZk+ku9us1+ZfHd8tiOxdQinkqm99Drff5eNNH6kLdgMGI4DRkkJufgD6wfnPpi8o7NdejoDV02aN2/eTbm5uX9mZsXMAgAlEom9APYSUXvXfbhxuvX/XcnMrYklhAjGYrE3XnrppZuIyHZ+PyoTdCWdCRMmnF5aWvoPv99f7EpFTriDcBhrbUNDw+JUKvUPr9f70b59+5oSiYSdl5cX8Pv9BbZtn5OdnX2F3+8fKaWEw1gEMyuPxyObm5s3r1+/fv7GjRu3oAutnlz6Jk6cOGvYsGFP+v3+grb0maYplFJIJpOb6uvrnxJCvKiU2ldfXx+zLIuLi4sDHo8nL5lMluXk5Fzu8/nGG4ZxCH2GYchUKpXat2/fZ1599dXHu9L92k1C/udV3q/MP8X8rc+EtiwnYVlDmyYEJNAS49iG/Wr5riZ+ckQe1r2zRzes3EXxMQXsmTZA5NU005CCIK4eUyjP9vuFaadDA4TSUB4/yZo6teHGaGzm01txwIle7/QLjp3CfcMZWf/8UuDfY0vk6XY8HffIDDYlAAPUEuOW9/eqJc02Hs31Yus/1uvabbWUGlXM3otHiII9zVxakEnlo/Pl3EBQ+GCxthSICASGMoIkV2+1npz0x8RVzLDddNXOrmFL2F/i0XKboHQYkpSgAzFdG7fRJAiCCUpoCE4HPANI18xzut6wEwWQ9uqkmw3aRMjaegDLTvtzbAGHITqUsBx3Kf82F1l/uND/8JAiOd46yKy0YYBaUkj8fXXyqntWYVtXGER3uvIAgJuSQ+WJH797s3/cpCHyGsuxpbGG0AR1+lDj20+VY9ullck/H6lYWi+AACAQCBQT0VGbHBwJTkUDt1OyB11okunGR0Wj0TeEEJcOGTLk0WAwOCyVSikAQinFALRpmvlFRUXlzFyeSqXg8/matNYpKWWmz+fzmKYJrTWUUmxZFjvqr+3xeIyWlpYPt2/f/qmNGzdu6QozcOlzmNarRBQaOnToI36/v5/DtITDeOD3+0cEg8Fv2Lb9jUQiobKzsxsdhpvp9Xrbpc9VA5PJZNPOnTtvfuutt7rErICDITRUnvzd368Czh1h/jbTB1hJ2ELAsGywUNBBDwWmDDbnTZGYhwQjy8vWaSXcbAr4B2aSTwwVabnGZthJVkSQSkN7PJCNMR1/6gN90zPbUI8KCOqiikUAh++A+MEiNP5gWezyH871Vw4fKGfrBGylIG0FgoYKeih4+jDzYgi+ONHCyAuIA0pRwiPZ3y+TsicOa0NjIs3wiKENCQ0Tcv02++lz/5i4rivMqkOa02FOyA2I/FyB/KNeYXtgAF5Cc0oVuF91rBI6ssEXJ3kLs310ZjoKCsQEkIDSgsyVH6W+/IXn7OVHs1v1KsrdltrxhZtvC5xSWiyn2Yn0A8IKOpgpzMG54jQAf8a43nWpOtJQKwDAYQg9VQe1bdsS3aiMCbSmpsj33nvvrf379583ceLE32VmZs4nIrdJBimltNtpxzAMMk0z060bpbXWqVSqrTTIRCQMwzAaGxtf27Bhwxc2bdpUDUB0hRm4cGxd8r333luaTCYvKi0t/WNWVtYMpRSYWTMzuc08iIh8Pp8UQuQC6WKJzMypVKptSwYmIuH1emVTU9OmXbt2feXtt9/+d1eZlQsqP8i0HrqMd58z0vjf4nw5UCeYCdCaQVqBSTEzwFIAA7LJBFEuGICGsmPMSDciJYdO5fGTbGzW8dc/sj77hX+lllf2QMU62NUHex9dEz/v3Vv8fxjXT95gegi2xYoZZCkwVJpmnyQxLI9yHHcHtA1tx1g7DgACQMTQhpdE3Ga8t8X+/cx7E18lgqqg7jdLZQa3tcVoBe52qjZBSQGp9cEROvQqEIGjlRA/WpLcvOQjtaC+hZVTv9qSfjLf22ndOeeB5H0chtEVZtVeblFXuAoBjAoQEVqWbNILahv0LsNDknWaro92quW3PBO/jcMQTtBqb0JKKUkIIaWUJKUkwzCEYRiy7aerEEKYjurWk4oCKhQKyZ07d2587rnnLtu1a9c34/H4DsMwhGmawqk9zwBYa62VUsq2baWUapvRzc7x0rKs+N69e+985ZVXLtm0aVO1E37Qk/VUoVBIrl+//u1XXnnlwpqaml+kUqkW0zSlaZqiTbkc1lor27aVbdvKMaq70etoQ19q//79f3777bfP7gmzcq/b9QZ+5p+pJ36yNDZ79Vb7bwmLyfBAmh4SMl1CRgOAUqCUBZ1KQacs6JRKlzxihjYl2PSSMDyQW2vVmqfXqQvOfyT1BPeCiSISaU1FS5x6V/zGF9fbn9vdqLcbXkjTS8KUYKQbY5ClgVQKOpVM06m0ozUxtCkA00vC8EJsrVebXlqf/PTMexNfIYLiI5h1joYAANOAMAyQefAjTLObHwlDGBDeNtXlj2h0Ly+HSnsIky88c7W4/dwR4teeoDA37bSfnXJX4hvMEJ01HvYmKJKOBKZoYmt+pvcz55Qaz2ZmCv+eerX9J2/FrnpzB+I4Bp7CZDIZb2lpSSHdcimdcpUOJ+C23kJHnXL/TkdMduBNdKU2y7K8qVSqFl1QBw+Hs2EFgMSrr776q4EDB1aOHDnyS5mZmZdIKcf7/X7p0NL6ISI4bcNg2zbi8Xh9Mpn899atW/93/fr1bzhDd0uyao8+Rz2sffnll/9r5MiRDw4bNuxWr9d7ntfrLXFrsXdEn2VZiMViNZZlvbR9+/Y/rF27djnQaifrMX1uQj1FsOXXK+PX//gc467LR3m+XpBJMwv9VCIyKH13NNIRUUD67gpKr7oBxBo0knFet3KPevR3TyfufKoWTZVd7E14JKTrA4AQBlEk8dfwVDwze7z3K5OLjZDXwNhAjhCwkVb92hqEBaV3uwAa6jWsmF7z5i796MOvJP7w8DbUO3uZ0YOX0o44IG2dalOJwpXaD61h10Gbsna6ASmKk3d3I7euXaeEGzes4d1bfPcVZoo5t/8zNuPhDaitIFBnubFrdP/GTHNyZI7nrQwvGa7R/b1t9iundsLo/jG6HCP8q9f5vzhiAP36herU7Ouest/sTYO7AwLAp5xySgERjRdCKK21EkIopVTr9btMyQmUTN8oZmr7/eEDt0nW9Wut927YsOH93qDX6cTs0pZ75plnzsrIyBitlJoqpRxhGEYhERla67hlWbsAvCelfGfdunVvb9269V2glRH0hsr7MfocZsQAkJ2dXTpt2rSZRDQOwGTDMAYZhpGDtCrbmEqltgsh3gOwdt26dSt37NixBoDbyIPQy9U3nSYqrQbgUdkY+tBV3nO9hhjTnMT4bC8GZ3mRZ0jypBRbB+LY35ziD/OC9O5Htbz6/IcTzwBoBpweB737LLbiMFNM8LnP+C4alkun1scwOTdAI7I9yIfjpauPc119kjcU+unttbVq5RWPpJ4D0hkmvWXSCZfBJ1jOhCatCawJbAgIaBauZiWc79s7XzjPg227f4MB9h1IiJr/XWG9i84qY464S6GxyLhhAkqBrhvO3YDOb8w0Jzd9J2hxRQanvp9h848y+d2b/cscArsMZ1zxvckY49La9VH+Y0EdRJIHR48enT916tSCwYMH5wLwtv3RSZk5HlUwRDv0GXl5eVlTp04tmDp1akFpaWk2gEPahB0v+sJhiHayJ4wZech66GIUPHFtsOjui1EwKh+ZwMHjBDnVR4/Ds8hIl2k+bO8YoVJkPxvKKHzi2mDRE1cEiy4bgpy2NMrjSGNvolvEptMDuvbWPZKE1dmwho7g+v+PsVcQSDOAY7ZRnFiuY0E/lZWVyaKiIn788cdVe0UImZnmzJkjq6qqGMe/8akoKysTc+bM0YsWLdLt0RcOh8XSpUvFkRp1HCuEnWa+qICWAu12jxQEqMcgl64FzU0b1o93XB4tCUPOGQeWV0Edkcbfg5zKC71OY29VkGiLtWPB3aqiwYfrol3AsZKwDh+/D0cFdfA5WXDS0+fuA+aD/4+TkUZu8zn5aOwWupQ0eSzSXXoLJ1EnnZMdJ+09dHDS03fQu3IiyTginKjkE01G76NPKulDH/rwicFxZ1i93aesD33ow/8dnCgJSx/y4RNferUPJxT/5yrG9qF7OO4MizSE14AHJgnTIBMeEkQcPN509OHkgBOjpp2YLBkOh8XJUNm1DycnjjvDSjHFm1O8LhHTG5sTekOiRW+M29gBnPzW1j70PoiIMzIyCtzUzEgk4jKvPvThYzgRbzI36M99KN1I5T4v3/8dtEa6l5WV3eD3+38MYNkHH3zwm6KiojEZGRnGSy+9dJcTyd7HvPrQih4VsO8mtKBDmRMD7bfBPblBSKfAtH7RhTSWI8XEdGaMQ+Z2Kmv2JIWmNTG6m+dT23LJnVgHIiI9f/78PMMwbtVarzdN84wJEyaUMzMaGhpuA4CKiopO1QTrJg196EOncbIGBXYWRyqNfMKupyddcrqLjjrfdLLhrCgtLR3v/H/+BRdc8N9nn33259oO0xMaetiVpw8nIT6JzOJEo/WtP2HChBKv1zs8lUpl+Xy+bZs3b95SW1vbdLQBTj/99DwiyiQiy+2OLITgZDIpY7HYgdWrV7cc6fzBgwfnDhw4MMuyLJ/P52OPx9Py8ssvNwI46tyHY+rUqSaA/qZpNr/xxht1XTjV7ZGKESNGFObm5o6xLKuYiHa0tLRs/OCDD/Z3hY6ysrKCqqqqRgCprtAA515MnDixKBAIjEwmk/kej+ejN998czOAlrbH9OGTjxOhEn5i4dpUsrOzcyZNmvTD3Nzcm7xer88p4ofCwsKNtbW1n+2oSYN7vtfr/WNubu4lOp2V5rZ7UkIIb01NzQ8BRNqp70QAuKSkxD9y5Mh/ZmVlTQcghBBaKaUvuOCCrbFY7LdVVVV3oxPqnTO+9vv95+Xn5z+aSqXeAHAR0gyjU2otM/OsWbO+lJub+71AINDfKQaIVCq1b8CAAV9cunTpE0epU+U2nBidkZHx3AUXXNBUU1Mzb9WqVbvROUZDAMSsWbO+mZeX9y3DMLKZ2SYi44ILLli9e/fur7/77rtVfbaw/xz0icxdh5w+ffpfhgwZ8mWl1OotW7Z8bevWrdfs2LEjLKUUwWBwGHCw20tbCCEYAEzT7GcYhq+hoeHNxsbGqoaGhuWNjY2v19fXvxmPxzcDHTe18Hg8noyMjMlKqb1btmz55rZt227btWvXL03TNAYMGPCHc8455ytAa0v7DuF0pWGfz3dFTk5O0Ov1zpoyZcr4zpzrqFp61qxZt5WUlPxeCJH48MMPw5s2bbr2o48++rpt25uDwWARANTU1HQoxZeVlQkAuqSk5ILs7OzBWVlZ4zIyMqYC6a7OnaFh9uzZ3xw0aNBP4/H4sm3btp1z4MCB03bt2vUVn883bMiQIZUjR44c4zCrvme9D/934G7iM8444zPXXnstz58/vxKA/7DDgoWFhRkdjeHadebNm7f84osv3tHO+UcCAUBhYWHGpZde2nzRRRc9ftjvBZdeeunWiy++eE9paWn20VrKOwhcccUVm84999yayy67zD7vvPNuQ8claQ6hY/Dgwf0vv/zyuksvvXRLXl5eSTvHHdWe5tJ4/vnnv3rJJZd8dOmllzacf/759wFHtccJAJg6dWr/K6+8su6iiy56EYcxpHHjxp0RCoV43rx59wB99qz/FPTdxE7CVWtyc3MXtrS0JDZv3vzfRBQvKyszwuGwcDZEy759+5qPNhYRCSLS+fn5ree2KfR3VLsiMwsiIudcY+HChSaA/clk8uFgMFhcUlJSxMzoqNGpywymTZs23TCM0h07dnxXa/2OlDIEgB9//PEOMw9cyWfo0KFX+P3+3P379/+srq5ux8KFC033WhxGdLTsBcHMGDly5ACfzzezoaHhL01NTU/5fL5zAWREo9G2NdwPp4EAIDMz82zTNHNra2t/SUTapaGyslKuXbv2tVgs9pzf71/Qv3//gKOe99lsP+Hos2F1DgSAg8FgkcfjmWZZ1lubN2/e5Nip7KqqqrbHAZ1pjcQsbduW1dXVVFNTQ+Xl5a11+490ntbarWQqnY40WLVqlQQgg8HgTiLSeXl5ndqYBQUF8y3L0rW1tY8MHjx4dDAYvKVfv36Fe/bs2Yej2JCys7PPsCzL3rlz5/MA6J577lFHo70tysrKRFVVlS4oKJhvmqbYtm3bP4YPHz7CNM1rzzrrrImvvPLKa0er024YxrxEItG4Zs2aVcwMlwa3k7RlWcv9fv8Fubm5p+zevfvdo11TH05+9ElYnYArVQwaNGiMlDIQj8ffBkDt2ak6CyfeqDkajaqqqirbbb+Oo2woIQQ7PRBlOBwW1dXVNHXqVABQLS0tpZZlifr6+tgRhmht3y6EKLcs6+19+/Y1x2KxF03TDIwePXo+AJSVlbWnkpEjfRGAYZZlNWzdunUHuhHDNWfOHA0Aubm5FyYSiQMfffTR6h07dqyyLCvu8XjOA9DanfpwuN8HAoHhQojtzc3N+9wUHwAoKipiANzY2LiDmdnj8YwADkpmffjkok/C6gIKCwv7GYaBZDL5IdrfoJ3atMwc9/l8OXPnzo14vd4mrbWHiLQQIpVKpe5dvHhxLTqQBoQQ7HSRUW28kCovL68kEAhck0gkVlZVVe0DQB10ZyYAPGLEiOE+n2/4gQMHfgIAdXV17+bl5TX6fL75AB6aM2eObiM5Hg6f1rpIa12L7lUnpUgkoocMGZIjpZxn2/bTAHjDhg1bR40atcbv918CILxo0aIjrqdpmh5mTgDtB5kmk8m9ACgjI2NAN2jsw0mIPgmrE3A9XaZpmgCglGq3s03//v0DI0aMyOpoHEedAzMnvV5vIDs7+1s+n+9/gsFgxO/3/8Dj8fyYiAYA6ND+pLUmIooJISaceuqp10ycOPGSGTNmfO2ss85aKqUsrqur+xGApCMVfmzDO545DB06dJ4Qgurq6p4EgOrq6r1a68WmaV48ZMgQX0c2HyfnzyQir9Y60d4cR4MrsRYXF0/2er3ZiUTiBXf4lpaWFwzDGD18+PARzlxdfkZdCcy27QMAQERe4Mgeyz58MtAnYXUBlmVZzAwhhKft927M1bBhw+7Izs6esWnTprOPNI6U0tvU1FT32muvLWDmuNOXUCqltM/n2wS0ppZ8DHl5eaS1jgeDwdLBgwc/YBiG8Hg8iMVim3ft2vXZ119//R9O3FG79iRXcvL5fOcnk8mdzc3Nq+FIXQcOHFjSr1+/KwoKCqZt3br1VXRs87EBJIQQQXShbf3hyMnJuVgpZa1fv36p+10ikXg6Ly/vu6WlpWWbN2/eHAqFKBqNdjgGHSGk3jCMHKdZbLdbp/Xh5EKfhNUJODYR1NXV1di2Db/fX9r2d9eW5fV6B/r9/uFtfjpkM7lxWMzsYeaG2trapXV1dW/s37//lb179y7dv3//sh07dsSdw9tlWAcOHNBElNnc3LyipaVlWkNDw5cBoKGhoeq111570PHSdWRTokgkoqdOnZotpTwPwNObNm1KLly40ACA3bt3P8fMKCoqOh84KI0dMkCaP8SJaJ9hGLnoRPjC4UNUVlZqAKbH4zkvlUqt3r179zY31Wn9+vVrk8nkFsMwLgTAzrHtQmtta60DALBo0aLW49z7YZrmQCLi5ubmPcDB+9iHTy76GFYn4AZxfvTRR9Va6xa/3z+5veMc21K87VcdDCmJyASQifQ9kDhYxaIzMIUQHyxevPidxYsX/6Guru6f/fr1u3ry5MnTHFWuo9w6AgAp5SyPx+Orr6+vAoANGzYwAGzevPnDRCLxvsPMjKVLl7oG9tbrWbBggUS6VfwWIUR2v379BqBr4QJERDx16tThXq93nFLqSQAoLy8XoVBI1tbWNjFzldfrnQcgy2G+7Y4fj8f32LZdEggE+rUNCXFUP8rJySlhZkokEhuAg6piHz656GNYnYNmZmppadmbSqVWejye6WPHjh37+OOPq6lTp5pbtmwRAAQRSafd/NHAzGxkZWUZoVCIysrKKBQKUSgUok4GOJIQgkOhkGRm2rNnz/csy/L079//pzgo8Xxsk7uSR2Fh4QW2bdu7du1aDABVVVXKic1SzPwPj8czZcaMGf2IiDuypdXV1b3u8/lkaWnpXAA8duxYEw7TdcY6UgwVZWZmniOlRH19/fJ2xn7JMIzMGTNmnOWcc8iauNcRj8dXBoPBjLFjx052mKDT2xgAwIZhnGlZVpMQ4gOgYzW7D58c9DGsTqK8vFwAQF1d3V0+n89bUlLyC2bOXrVqlbVq1SoLgLZtW7mG9SOBiDQzc2NjY+PhYQ2H5x8e4XwZjUbV0qVL5erVq9fU1dXdmZOTc/asWbPKAeh2UlvckAQvEZ2fSCRWbNmypQaHMZY9e/YsJSIRDAYvBIClS5ceMo4rpezcufOJeDxen5eXdzuArOrq6hScumZHCs9wU4ICgcAlqVRq75tvvrkKAKLRqHbHrq2tXaKUsgoKCuYdaR1aWlqeV0rpgoKCbzCzcO9DVVWVPWbMmLOCweC8RCLxSHV1dXNfpPt/BvpuYicRjUZVOBwW+/bte7KmpuaJoqKiCy666KKXp0+f/s0xY8ZcOHHixE9lZWWdKYRw1TugAylDa235fL78OXPm3HP22Wf/Zu7cuXedffbZfzj33HPvnTNnzs3Awbb37cExNEsAWLp0KQDQ+vXrfxaLxXYUFhb+ePDgwblOpHjr/Q2Hw8TMmDVr1qhgMDhca/000Bpvxc7xePvtt1+3bbshMzPzEuBgvJQLJ1hVfPjhh3tra2u/n5GRMe6SSy556dRTT71++PDhZ4wePXre9OnTv3H66aef6czb9hkTkUhEFxYW9vN6vdMTicRyAE3OMRyJRJiIsH79+hrLspYbhnEpAI/DAFvXw70XK1eufP/AgQMPFRYWnjtv3rxHhg4dOnvQoEHjpk6d+vlRo0Y9Zdv23m3btv0Cadsd0Bc0+olHn5ewC3BUiuSmTZs+d84553wUDAZvLS0t/YXWGkop2LZtNzc3/wUHN0a7GySZTDYD8GRnZ9/gOrmYGaZpoqWlpQTA3R0VrzNNUyeTyRQzu2VYdCgUEtFodF9tbe238/LyHhgyZMgt27Zt+x/newAH1ShmvigWi1k1NTVLgINew/RPTEQUb2lpeVpKeVG/fv0KI5HIPhzmCYxEIu6xv589e7bMzs6OjBw58n5mtplZEhHV1tZ+B8ByR0LTQFodjEajGDNmzARmzmtpaalsSxsAnj17tlFVVWUnEonn/X7/T8aNGzd87dq163CYx9K5F/Zbb7112/Tp01VmZubnpk6dugDpahO+RCKxqaam5pZ169ZtPJLXtA+fLPTFpXQdrRtn3LhxY0pKSkpt287wer0NO3bs2Lt69ep1ABJHGmDq1KnD/X5/P6VUyjAMAgDbtpmIDK31vjfeeGMTOg4poJkzZ55qGMaBV1555cM2xxEAcdppp52ak5PT/O9//3tDe3SfdtppI5VSOStXrnwbHw/6JAA8ceLEomAwOIKZ33njjTfiR6IFAJ9yyinDBg8ePLm+vv4Uv9+/0zCMXfv371+zZs2avYedSwB41qxZualUarxSatWqVati7R0zYMCA/IEDB55SW1tbvWXLloYOaGj97vTTT58SDAaHaa29pmnue+GFF94BsL+vtEwf+pAux9uTSpv/MWhTVcF3xhln3D9p0qRTjzMJHYZi9dmt/vPwf2t39TLaK4Hi2IKO9kYXHeW1jR07lo9meHeSghkfD9gkRw1s77dD5j5SUnGbcTqTdkPhcJgikUjwsssue5eIfmGa5j01NTVUVVXVUcCmO/6R1qozxxxyLJBev+rqajrKGvShD334PwoJAGPHjh1x6aWXNl988cWXAyemvnwf/vPRJ2H1oSdwnx/vxRdf/CQzD/nwww9nrVu3ro6Zu52y04c+dIQ+Hb8PPQUDsBOJxLvbt2+/obq6us6JRetjVn3oQx9OevRJ7X04Zvj/wfTW2BTnFlQAAAAASUVORK5CYII="
    def file_generate(self):
        self.gets_entry_values()
        self.pdf.create_file()

    def gets_radiobutton_values(self):
        aux = self.troca_aparelhos_var.get()
        if aux == 1:
            return ''
        if aux == 2:
            return 'e troca de aparelhos de ar condicionados não eficientes por modelos com selo PROCEL'
        if aux == 3:
            return 'e troca de aparelhos de refrigeração não eficientes por modelos com selo PROCEL'

    def gets_entry_values(self):
        self.pdf.name= self.document_name_entry.get()
        # self.pdf.client = self.document_client_entry.get()
        self.pdf.project_atributes['Cliente'] = self.client_name_entry.get()
        self.pdf.project_client['Nome Cliente'] = self.client_name_entry.get()
        self.pdf.project_client['Nome Fantasia'] = self.client_name_f_entry.get()
        self.pdf.project_client['CNPJ'] = self.client_cnpj_entry.get()
        self.pdf.project_client['Num Cliente'] = self.client_num_entry.get()
        self.pdf.project_client['Modalidade da Tarifa'] = self.client_mod_entry.get()
        self.pdf.project_client['Classe/Subclasse'] = self.client_class_entry.get()
        self.pdf.project_client['Endereco'] = self.client_address_entry.get()
        self.pdf.project_client['Cidade'] = self.client_city_entry.get()
        self.pdf.project_client['Estado'] = self.client_state_entry.get()
        self.pdf.project_client['Telefone'] = self.client_phone_entry.get()
        self.pdf.project_client['E-mail'] = self.client_mail_entry.get()
        self.pdf.project_client['Contato'] = self.client_contact_entry.get()
        self.pdf.project_client['Ramo de Atividade'] = self.client_activity_entry.get()
        self.pdf.project_atributes['Usos Finais'] = self.final_use_entry.get()
        self.pdf.project_atributes['Energia Economizada'] = self.energy_economy_entry.get()
        self.pdf.project_atributes['Demanda na Ponta'] = self.demanda_entry.get()
        self.pdf.project_atributes['RCB'] = self.rcb_entry.get()

        self.pdf.project_atributes['Trocas'] = self.gets_radiobutton_values()
        
        print(self.pdf.project_atributes['Trocas'])
        print(self.pdf.name)
        print(self.pdf.project_atributes['Cliente'])
        # print(self.pdf.project_atributes)

#building the application
class Application(Application_Functions):
    def __init__(self):
        self.root=Tk()
        self.pdf = Relatorio()
        self.base64_images()
        self.screen()
        self.window_buttons()
        self.frame_logo()
        self.main_frame()
        self.pages()
        self.widgets_page1()
        self.widgets_page2()
        self.widgets_page3()
        self.labels_frame_logo()
        self.root.mainloop()

    def screen(self):
        self.root.title('ECOSOL - PDF AUTO GENERATE')
        self.root.geometry("1000x700")
        self.root.resizable(True,True)
        #self.root.maxsize(width=,height)
        #self.root.minsize(width=,height=)
    #set first frame definitions
    def frame_logo(self):
        self.fr_logo=Frame(self.root,bd=4,)
        self.fr_logo.place(relx=0.3,rely=0.01,relwidth=0.4,relheight=0.2)
    #setting main frame definitions
    def main_frame(self):
        self.main_fr=Frame(self.root,bd=4,highlightbackground='black',highlightthickness=0.5)
        self.main_fr.place(relx=0.01,rely=0.2,relwidth=0.98,relheight=0.725)
    #setting notebook in main frame
    def pages(self):
        self.notebook = ttk.Notebook(self.main_fr)
        self.page1 = Frame(self.notebook)
        self.page2 = Frame(self.notebook)
        self.page3 = Frame(self.notebook)

        self.page1.configure(background='lightgray')
        self.page2.configure(background='lightgray')
        self.page3.configure(background = 'lightgray')

        self.notebook.add(self.page1,text='Document Settings')
        self.notebook.add(self.page2, text='Client Information')
        self.notebook.add(self.page3, text = 'Resumo do Projeto')

        self.notebook.place(relx=0,rely=0,relwidth=1,relheight=1)
    #setting logo
    def labels_frame_logo(self):
        self.ecosol_img = PhotoImage(data=base64.b64decode(self.ecosol_logo))
        # self.ecosol_img = self.resize_image(self.ecosol_logo,2,2)
        # self.ecosol_img_aux = PhotoImage(self.ecosol_img)
        self.ecosol_img.subsample(2,2)
        self.ecosol_label_img = Label(self.fr_logo, image=self.ecosol_img)
        self.ecosol_label_img.place(relx=0.0,rely=0.0,relheight=1,relwidth=1)
    def widgets_page1(self):
        #creating page1 labels
        self.document_name_label = Label(self.page1,text="Document Name:",bg = 'lightgray',)
        self.document_client_label = Label(self.page1,text="Client Name:",bg = "lightgray")
        
        #locating page1 labels
        self.document_name_label.place(relx=0,rely=0,relwidth=0.12,relheight=0.1)
        self.document_client_label.place(relx=0,rely=0.11,relwidth=0.1,relheight=0.1)
        
        #creating page1 entrys
        self.document_name_entry = Entry(self.page1)
        self.document_name_entry.place(relx=0.01,rely=0.071,relwidth=0.15,relheight=0.06)
        
        self.document_client_entry = Entry(self.page1)
        self.document_client_entry.place(relx=0.01,rely=0.181,relwidth=0.15,relheight=0.06)

    def widgets_page2(self):
        self.client_name_label = Label(self.page2,text='Nome do Cliente',bg='lightgray')
        self.client_name_label.place(relx=0.01,rely=0.0,relwidth=0.10,relheight=0.1)

        self.client_name_entry = Entry(self.page2)
        self.client_name_entry.place(relx=0.01,rely=0.08,relwidth=0.15,relheight=0.06)

        self.client_name_f_label = Label(self.page2,text='Nome Fantasia',bg='lightgray')
        self.client_name_f_label.place(relx=0.23,rely=0.0,relwidth=0.15,relheight=0.1)

        self.client_name_f_entry = Entry(self.page2)
        self.client_name_f_entry.place(relx=0.25,rely=0.08,relwidth=0.15,relheight=0.06)
        
        self.client_cnpj_label = Label(self.page2,text='CNPJ',bg='lightgray')
        self.client_cnpj_label.place(relx=0.45,rely=0.0,relwidth=0.1,relheight=0.1)

        self.client_cnpj_entry = Entry(self.page2)
        self.client_cnpj_entry.place(relx=0.48,rely=0.08,relwidth=0.15,relheight=0.06)
        
        self.client_num_label = Label(self.page2,text='N°s do Cliente',bg='lightgray')
        self.client_num_label.place(relx=0.7,rely=0.0,relwidth=0.15,relheight=0.1)

        self.client_num_entry = Entry(self.page2)
        self.client_num_entry.place(relx=0.73,rely=0.08,relwidth=0.15,relheight=0.06)

        self.client_mod_label = Label(self.page2,text='Modalidade Tarifária',bg='lightgray')
        self.client_mod_label.place(relx=0.015,rely=0.15,relwidth=0.115,relheight=0.1)

        self.client_mod_entry = Entry(self.page2)
        self.client_mod_entry.place(relx=0.01,rely=0.23,relwidth=0.15,relheight=0.06)
        
        self.client_class_label = Label(self.page2,text='Classe/Subclasse',bg='lightgray')
        self.client_class_label.place(relx=0.25,rely=0.15,relwidth=0.11,relheight=0.1)

        self.client_class_entry = Entry(self.page2)
        self.client_class_entry.place(relx=0.25,rely=0.23,relwidth=0.15,relheight=0.06)

        self.client_address_label = Label(self.page2,text='Endedreço',bg='lightgray')
        self.client_address_label.place(relx=0.463,rely=0.15,relwidth=0.11,relheight=0.1)

        self.client_address_entry = Entry(self.page2)
        self.client_address_entry.place(relx=0.48,rely=0.23,relwidth=0.15,relheight=0.06)

        self.client_city_label = Label(self.page2,text='Cidade',bg='lightgray')
        self.client_city_label.place(relx=0.705,rely=0.15,relwidth=0.11,relheight=0.1)

        self.client_city_entry = Entry(self.page2)
        self.client_city_entry.place(relx=0.73,rely=0.23,relwidth=0.15,relheight=0.06)

        self.client_state_label = Label(self.page2,text='Estado',bg='lightgray')
        self.client_state_label.place(relx=0,rely=0.3,relwidth=0.08,relheight=0.1)

        self.client_state_entry = Entry(self.page2)
        self.client_state_entry.place(relx=0.015,rely=0.38,relwidth=0.15,relheight=0.06)

        self.client_phone_label = Label(self.page2,text='Telefone',bg='lightgray')
        self.client_phone_label.place(relx=0.22,rely=0.3,relwidth=0.12,relheight=0.1)

        self.client_phone_entry = Entry(self.page2)
        self.client_phone_entry.place(relx=0.25,rely=0.38,relwidth=0.15,relheight=0.06)

        self.client_mail_label = Label(self.page2,text='E-mail',bg='lightgray')
        self.client_mail_label.place(relx=0.47,rely=0.3,relwidth=0.08,relheight=0.1)

        self.client_mail_entry = Entry(self.page2)
        self.client_mail_entry.place(relx=0.48,rely=0.38,relwidth=0.15,relheight=0.06)

        self.client_contact_label = Label(self.page2,text='Contato',bg='lightgray')
        self.client_contact_label.place(relx=0.72,rely=0.3,relwidth=0.08,relheight=0.1)

        self.client_contact_entry = Entry(self.page2)
        self.client_contact_entry.place(relx=0.73,rely=0.38,relwidth=0.15,relheight=0.06)

        self.client_activity_label = Label(self.page2,text='Ramo de Atividade',bg='lightgray')
        self.client_activity_label.place(relx=0.01,rely=0.475,relwidth=0.13,relheight=0.1)

        self.client_activity_entry = Entry(self.page2)
        self.client_activity_entry.place(relx=0.015,rely=0.57,relwidth=0.15,relheight=0.06)

    def widgets_page3(self):
        self.final_use_label = Label(self.page3,text='Usos Finais do Projeto:', bg='lightgray')
        self.final_use_label.place(relx=0.01,rely=0.0,relwidth=0.125,relheight=0.1)

        # self.final_use_entry = Text(self.page3,height=10,width=80)
        # self.final_use_entry.place(relx=0.01,rely=0.08)

        self.final_use_entry = Entry(self.page3)
        self.final_use_entry.place(relx=0.01,rely=0.08,relwidth=0.172,relheight=0.06)

        self.energy_economy_label = Label(self.page3,text='Energia Economizada (em MWh/ano)',bg='lightgray')
        self.energy_economy_label.place(relx=0.23,rely=0.0,relwidth=0.205,relheight=0.1)

        self.energy_economy_entry = Entry(self.page3)
        self.energy_economy_entry.place(relx=0.23,rely=0.08,relwidth=0.205,relheight=0.06)

        self.demanda_label = Label(self.page3,text='Demanda na Ponta(em kW)',bg='lightgray')
        self.demanda_label.place(relx=0.5,rely=0.0,relwidth=0.15,relheight=0.1)

        self.demanda_entry = Entry(self.page3)
        self.demanda_entry.place(relx=0.5,rely=0.08,relwidth=0.15,relheight=0.06)

        self.troca_aparelhos_label = Label(self.page3,text='Haverá troca de aparelhos?', bg='lightgray')
        self.troca_aparelhos_label.place(relx=0.72,rely=0.0,relwidth=0.15,relheight=0.1)

        self.rcb_label = Label(self.page3,text = 'Relação Custo-Benefício (RCB)',bg='lightgray')
        self.rcb_label.place(relx=0.01,rely=0.145,relwidth=0.17,relheight=0.1)

        self.rcb_entry = Entry(self.page3)
        self.rcb_entry.place(relx=0.01,rely=0.23,relwidth=0.172,relheight=0.06)

        #making a radio button to select the options to "troxa de aparelhos"
        self.troca_aparelhos_var = IntVar()
        
        self.op1 = Radiobutton(self.page3,text='Sem troca de aparelhos',variable=self.troca_aparelhos_var,value=1,bg='lightgray')
        self.op1.place(relx=0.72,rely=0.07,relwidth=0.15,relheight=0.1)
        
        self.op2 = Radiobutton(self.page3,text='troca de aparelhos de ar condicionados',variable=self.troca_aparelhos_var,value=2,bg='lightgray')
        self.op2.place(relx=0.725,rely=0.14,relwidth=0.23,relheight=0.1)

        self.op3 = Radiobutton(self.page3,text='troca de aparelhos de refrigeração',variable=self.troca_aparelhos_var,value=3,bg='lightgray')
        self.op3.place(relx=0.711,rely=0.21,relwidth=0.23,relheight=0.1)

    def window_buttons(self):
        self.generate_pdf_button = Button(self.root,text = "Gerar PDF",command = self.file_generate)
        self.generate_pdf_button.place(relx=0.45,rely=0.94,relwidth=0.1,relheight=0.04)

        self.quit_application_button = Button(self.root,text='Quit',command = self.root.destroy)
        self.quit_application_button.place(relx=0.85,rely=0.94,relwidth=0.1,relheight=0.04) 

Application()